from datasets import Dataset
import os
from dotenv import load_dotenv
import time
import traceback
import re 
from lightrag.utils import EmbeddingFunc
from lightrag import LightRAG,QueryParam
import lightrag.lightrag as lightrag_core
from lightrag.llm.ollama import ollama_model_complete, ollama_embed
from lightrag.llm.azure_openai import azure_openai_complete
from lightrag.kg.shared_storage import initialize_share_data, initialize_pipeline_status
import aiohttp
from ..server.server import mcp
import psycopg2
from azure.storage.blob import BlobServiceClient
from PyPDF2 import PdfReader
from docx import Document
import io
import numpy as np
from PIL import Image
import pytesseract
import fitz
import uuid
from typing import Optional, List, Dict
from lightrag.operate import get_keywords_from_query
# from pypdf import PdfReader
from lightrag import QueryParam
from fastmcp import Context
import asyncio
import json
import zipfile
import tempfile
from azure.storage.blob import BlobServiceClient, ContentSettings
from azure.storage.blob import generate_blob_sas, BlobSasPermissions
from datetime import datetime, timedelta
import base64
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import AnalyzeDocumentRequest
from crawl4ai import AsyncWebCrawler
from crawl4ai.async_configs import BrowserConfig, CrawlerRunConfig, CacheMode, DefaultMarkdownGenerator
from ..utils.azurecustomllm import AzureCustomLLM
from ..utils.access_validation import validate_user_workspace_access
from ..utils.request_context import request_var
from ..utils.db import db
# from tools.userManagementSystem import Session,UserMap
 
load_dotenv(os.path.abspath(os.path.join(os.getcwd(),'.env')))
 
azure_llm_api_key = os.getenv('AZURE_OPENAI_LLM_MODEL_API_KEY')
azure_llm_api_base = os.getenv('AZURE_OPENAI_LLM_MODEL_API_BASE')
azure_llm_api_version = os.getenv('AZURE_OPENAI_LLM_MODEL_API_VERSION')
azure_llm_deployement_name = os.getenv('AZURE_OPENAI_LLM_MODEL_LLM_MODEL')
 
azure_embedding_api_key = os.getenv('AZURE_OPENAI_EMBEDDING_MODEL_API_KEY')
azure_embedding_api_base = os.getenv('AZURE_OPENAI_EMBEDDING_MODEL_API_BASE')
azure_embedding_api_version = os.getenv('AZURE_OPENAI_EMBEDDING_MODEL_API_VERSION')
azure_embedding_deployement_name = os.getenv('AZURE_OPENAI_EMBEDDING_MODEL_EMBEDDING_MODEL')
 
 
os.environ["NEO4J_URI"] = os.getenv("NEO4J_DATABASE_NEO4J_BOLT_URI", "bolt://localhost:7687") or "bolt://localhost:7687"
os.environ["NEO4J_USERNAME"] = os.getenv("NEO4J_DATABASE_NEO4J_USER") or ""
os.environ["NEO4J_PASSWORD"] = os.getenv("NEO4J_DATABASE_NEO4J_PASSWORD") or ""
 
# os.environ["POSTGRES_HOST"] = os.getenv("POSTGRESQL_DATABASE_HOST") or ""
# os.environ["POSTGRES_HOST"] = "forgexpostgresql.postgres.database.azure.com"
# os.environ["POSTGRES_USER"] = os.getenv("POSTGRESQL_DATABASE_USER") or ""
# os.environ["POSTGRES_PASSWORD"] = os.getenv("POSTGRESQL_DATABASE_PASSWORD") or ""
# os.environ["POSTGRES_DATABASE"] = os.getenv("POSTGRESQL_DATABASE_DATABASE") or ""

os.environ["POSTGRES_HOST"] = os.getenv("LIGHTRAG_POSTGRESQL_DATABASE_HOST") or ""
os.environ["POSTGRES_USER"] = os.getenv("LIGHTRAG_POSTGRESQL_DATABASE_USER") or ""
os.environ["POSTGRES_PASSWORD"] = os.getenv("LIGHTRAG_POSTGRESQL_DATABASE_PASSWORD") or ""
os.environ["POSTGRES_DATABASE"] = os.getenv("LIGHTRAG_POSTGRESQL_DATABASE_DATABASE") or ""
 
embedding_dim = int(os.getenv("OLLAMA_MODEL_EMBEDDING_MODEL_DIMS", "1024"))
max_token_size = int(os.getenv("OLLAMA_MODEL_EMBEDDING_MODEL_MAX_TOKENS", "8192"))
base_url = os.getenv("OLLAMA_MODEL_BASE_URL")
 
pytesseract.pytesseract.tesseract_cmd = r'C:\Users\aalok\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'

async def llm_model_func(
    prompt, system_prompt=None, history_messages=[], **kwargs
) -> str:
    headers = {
        "Content-Type": "application/json",
        "api-key": azure_llm_api_key,
    }
    endpoint = f"{azure_llm_api_base}openai/deployments/{azure_llm_deployement_name}/chat/completions?api-version={azure_llm_api_version}"
 
    messages = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    if history_messages:
        messages.extend(history_messages)
    messages.append({"role": "user", "content": prompt})
 
    payload = {
        "messages": messages,
        "temperature": kwargs.get("temperature", 0),
        "top_p": kwargs.get("top_p", 1),
        "n": kwargs.get("n", 1),
    }
 
    async with aiohttp.ClientSession() as session:
        async with session.post(endpoint, headers=headers, json=payload) as response:
            if response.status != 200:
                raise ValueError(
                    f"Request failed with status {response.status}: {await response.text()}"
                )
            result = await response.json()
            return result["choices"][0]["message"]["content"]
 
async def embedding_func(texts: list[str]) -> np.ndarray:
    headers = {
        "Content-Type": "application/json",
        "api-key": azure_embedding_api_key,
    }

    endpoint = f"{azure_embedding_api_base}openai/deployments/{azure_embedding_deployement_name}/embeddings?api-version={azure_embedding_api_version}"
 
    payload = {"input": texts,"dimensions": embedding_dim}
 
    async with aiohttp.ClientSession() as session:
        async with session.post(endpoint, headers=headers, json=payload) as response:
            if response.status != 200:
                raise ValueError(
                    f"Request failed with status {response.status}: {await response.text()}"
                )
            result = await response.json()
            embeddings = [item["embedding"] for item in result["data"]]
            return np.array(embeddings)
            
def parse_references_from_response(response_text: str) -> List[Dict[str, str]]:
    """
    Parse the References section from LightRAG response text.
    
    Args:
        response_text: The full response text with References section
        
    Returns:
        List of dicts with citation_number and file_path
    """
    references = []
    
    # Find the References section
    if "### References" not in response_text and "## References" not in response_text:
        return references
    
    # Split by References section
    parts = response_text.split("### References")
    if len(parts) < 2:
        parts = response_text.split("## References")
        if len(parts) < 2:
            return references
    
    ref_section = parts[1].strip()
    
    # Parse each reference line (format: - [1] path/to/file.ext or [1] path/to/file.ext)
    import re
    lines = ref_section.split('\n')
    for line in lines:
        # Match patterns like "- [1] file.pdf" or "[1] file.pdf" or "- [1]: file.pdf"
        match = re.match(r'^\s*-?\s*\[(\d+)\]\s*:?\s*(.+?)$', line.strip())
        if match:
            citation_num = match.group(1)
            file_path = match.group(2).strip()
            references.append({
                'citation_number': f"[{citation_num}]",
                'file_path': file_path
            })
    
    return references


def generate_download_url_for_file(
    domain: str, 
    kb_name: str, 
    file_path: str, 
    workspace_id: Optional[str] = None,
    role_id: Optional[int] = None,
    expiry_days: int = 7
) -> Optional[str]:
    """
    Generate a SAS download URL for a file in blob storage.
    Searches in both main container and workspace container.
    
    Args:
        domain: Domain name 
        kb_name: Knowledge base name (may include workspace suffix)
        file_path: File path from LightRAG citation
        workspace_id: Numeric workspace ID (e.g., "199")
        role_id: User role ID (34 = SME, others = workspace user)
        expiry_days: Days until URL expires
        
    Returns:
        Download URL or None if error
    """
    try:
        connection_string = os.getenv('AZURE_BLOB_STORAGE_CONNECTION_STRING')
        main_container = os.getenv('AZURE_BLOB_STORAGE_CONTAINER_NAME')
        workspace_container = "workspace"
        
        if not connection_string:
            print(f"Missing Azure connection string for file: {file_path}")
            return None
        
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        
        file_path = file_path.strip().strip('"').strip("'")
        
        file_name = os.path.basename(file_path)
        if '.' not in file_name:
            print(f"Skipping non-file reference: {file_path} (no file extension)")
            return None
        
        valid_extensions = ['.pdf', '.docx', '.doc', '.txt', '.xlsx', '.xls', '.pptx', '.ppt', '.csv']
        if not any(file_name.lower().endswith(ext) for ext in valid_extensions):
            print(f"Skipping invalid file type: {file_path}")
            return None
        
        # Extract original kb_name without workspace suffix
        original_kb_name = kb_name.split('/')[0] if '/' in kb_name else kb_name
        
        # Extract knowledge_bases name if present in kb_name
        # kb_name might be in format "kb_name/knowledge_bases" or just "kb_name"
        kb_parts = kb_name.split('/')
        base_kb_name = kb_parts[0]
        knowledge_base = kb_parts[1] if len(kb_parts) > 1 else None
        
        # Build list of blob paths to try
        blob_paths_to_try = []
        
        # Construct blob path based on what file_path contains
        if file_path.startswith(domain):
            # File path already has full path from LightRAG
            blob_paths_to_try.append(file_path)
        elif '/' in file_path and not file_path.startswith(domain):
            # Partial path, prepend domain
            blob_paths_to_try.append(f"{domain}/{file_path}")
        else:
            # Just filename, construct full paths with knowledge_bases level
            if knowledge_base:
                # New 3-level hierarchy: domain/kb_name/knowledge_bases/[workspace_id]/filename
                if workspace_id and role_id != 34:
                    # Workspace user pattern
                    blob_paths_to_try.append(f"{domain}/{base_kb_name}/{knowledge_base}/{workspace_id}/{file_name}")
                # SME or no workspace_id pattern
                blob_paths_to_try.append(f"{domain}/{base_kb_name}/{knowledge_base}/{file_name}")
            
            # Fallback to old patterns (for backward compatibility)
            blob_paths_to_try.append(f"{domain}/{base_kb_name}/{file_name}")  # Old SME pattern
            if workspace_id:
                blob_paths_to_try.append(f"{domain}/{base_kb_name}/{workspace_id}/{file_name}")  # Old workspace pattern
        
        # Try searching in both containers
        containers_to_search = [
            (main_container, "main"),
            (workspace_container, "workspace")
        ]
        
        for container_name, container_type in containers_to_search:
            if not container_name:
                continue
                
            for blob_path in blob_paths_to_try:
                print(f"Attempting {container_type} container, blob path: {blob_path}")
                blob_client = blob_service_client.get_blob_client(container=container_name, blob=blob_path)
                
                try:
                    if blob_client.exists():
                        print(f"✓ File found in {container_type} container: {blob_path}")
                        
                        sas_token = generate_blob_sas(
                            account_name=blob_service_client.account_name,
                            container_name=container_name,
                            blob_name=blob_path,
                            account_key=blob_service_client.credential.account_key,
                            permission=BlobSasPermissions(read=True),
                            expiry=datetime.now() + timedelta(days=expiry_days),
                            content_disposition=f'attachment; filename="{file_name}"'
                        )
                        
                        download_url = f"https://{blob_service_client.account_name}.blob.core.windows.net/{container_name}/{blob_path}?{sas_token}"
                        print(f"Generated download URL for: {file_path}")
                        return download_url
                except Exception as e:
                    # Log but continue trying other paths
                    print(f"Error checking blob {blob_path}: {e}")
                    continue
        
        print(f"File not found in any container or path combination: {file_path}")
        return None
        
    except Exception as e:
        print(f"Error generating download URL for {file_path}: {e}")
        return None
async def initialize_rag(domain: Optional[str] = None, kb_name: Optional[str] = None) -> LightRAG:
    data_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'data'))
    lightrag_database = ''.join(char for char in f"{domain}{kb_name}" if char.isalpha())
    os.environ['NEO4J_DATABASE'] = lightrag_database
    print(''.join(char for char in f"{domain}{kb_name}" if char.isalpha()))
    rag = LightRAG(
            working_dir=data_dir,
            llm_model_func=llm_model_func,
            embedding_func=EmbeddingFunc(
                embedding_dim=embedding_dim,
                max_token_size=max_token_size,
                func=embedding_func
            ),
            graph_storage="Neo4JStorage",
            workspace = lightrag_database,
            vector_storage="PGVectorStorage",
            chunk_token_size=1000,
            chunk_overlap_token_size=200,
        )
    await rag.initialize_storages()
    initialize_share_data()
    await initialize_pipeline_status()
    return rag

@mcp.tool()
async def query_rag(
    domain: Optional[str] = None, 
    kb_name: Optional[str] = None, 
    knowledge_bases: Optional[list[str]] = None, 
    question: Optional[str] = None, 
    user_prompt: Optional[str] = None, 
    history: Optional[list] = None,
    mode: str = 'mix',
    workspace_id: Optional[str] = None,  # ADD THIS
    role_id: Optional[int] = None        # ADD THIS
) -> dict:
    """
    Query the RAG system with a question and optional user prompt.
    If knowledge_bases is provided, query each and aggregate results.
    """

    digit_map = {
            '0': 'zero', '1': 'one', '2': 'two', '3': 'three', '4': 'four',
            '5': 'five', '6': 'six', '7': 'seven', '8': 'eight', '9': 'nine'
        }

    if knowledge_bases is None:
        knowledge_bases = []

    # workspace_id may be missing in some tool invocations; guard against None.
    if workspace_id:
        result = []
        for c in str(workspace_id):
            if c.isalpha():
                result.append(c)
            elif c.isdigit():
                result.append(digit_map[c])
        workspace_id_alpha = ''.join(result)
        if workspace_id_alpha and workspace_id_alpha not in knowledge_bases:
            knowledge_bases.append(workspace_id_alpha)
            

    if history is None:
        history = []
    if not question:
        question = ""
    if question:
        question = question.strip()
    if not question.endswith("Please provide detailed insights from official documents and reports."):
        question += " Please provide detailed insights from official documents and reports."
    try:
        user_prompt = f"""---Role---

You are a helpful assistant responding to user queries about Knowledge Graph and Document Chunks provided in JSON format below.

---Goal---

Generate a concise, accurate response based on the provided Knowledge Base. Follow all Response Rules strictly. Use both the conversation history and the current query to guide your response. Do not include any information not present in the Knowledge Base or conversation history.

When handling relationships with timestamps:
1. Each relationship has a "created_at" timestamp indicating when we acquired this knowledge.
2. When encountering conflicting relationships, consider both semantic content and timestamp.
3. Do not automatically prefer the most recent relationship—use contextual judgment.
4. For time-specific queries, prioritize temporal information in the content before considering timestamps.

---Query to be answered---
{question}

---Conversation History---
{history}

---Response Rules---

- **Format**: Use multiple paragraphs with markdown formatting and section headings.
- **Language**: Always respond in English, regardless of the language used in the question.
- **Emphasis**: Highlight all referenced information using **bold text**.
- **Inline Citations**:
    - Cite the source **immediately after** the referenced information using square brackets (e.g., [1], [2]).
    - Every time a source is used, it must be cited—even if it has been cited earlier.

- **Reference Mapping**:
    - Assign each source file a unique reference number starting from 1, in the order of **first appearance** in the main answer.
    - Maintain a mapping between source file names and their assigned reference numbers.
    - Use this mapping consistently throughout the answer.

- **References Section**:
    - **CRITICAL: Only cite information that comes from actual document files, NOT from entity or relationship names.**
    - **Only include a "References" section if actual document sources were cited in the main answer.**
    - If no information was found or no sources were cited, do not include a "References" section.
    - When including references, list the **complete file path** (e.g., Industry/SubIndustry/filename.ext) that was cited in the main answer.
    - The file path MUST be exactly as it appears in the knowledge base storage system.
    - List them in ascending order of their citation number.
    - Each file path should appear **only once**.
    - Do not include any file name that was not cited in the main answer.
    - Ensure that every reference number used in the main answer appears exactly once in the "References" section.
    - **DO NOT cite entity names, URLs, or server names - only cite actual document file paths with valid extensions.**
    -  Ignore any file whose content is generic, such as "This text file belongs to Hagerty area in Demo Instances domain" or similar boilerplate.
    **Reference Format Examples:**
    
    ✅ CORRECT (actual files with extensions):
    - [1] Banking/Asset Management/Portfolio_Analysis.pdf
    - [2] Finance/Reports/Client_Report.docx
    - [3] Insurance/Policies/Coverage_Details.txt
    
    ❌ INCORRECT (entity names, not files):
    - [1] Demo Industry
    - [2] ForgeX-Dev-KB.AzureWebsites.Net
    - [3] Company Name
    - [4] Server Configuration

- **Integrity**:
    - Only cite references that are actual document files from which the information was extracted.

- **Note**:
    - Always give answer never mention Sorry, I'm not able to provide an answer to that question or no-context

    """

# - **Continuity**: Maintain logical flow with the conversation history.
# - **Ignore**: - Ignore any file whose content is generic, such as "This text file belongs to Hagerty area in Demo Instances domain" or similar boilerplate.

# - If the answer is unknown, say so clearly.
#         user_prompt = f"""---Role---

# You are a helpful assistant responding to user queries about Knowledge Graph and Document Chunks provided in JSON format below.

# ---Goal---

# Generate a concise, accurate response based on the provided Knowledge Base. Follow all Response Rules strictly. Use both the conversation history and the current query to guide your response. Do not include any information not present in the Knowledge Base or conversation history.

# When handling relationships with timestamps:
# 1. Each relationship has a "created_at" timestamp indicating when we acquired this knowledge.
# 2. When encountering conflicting relationships, consider both semantic content and timestamp.
# 3. Do not automatically prefer the most recent relationship—use contextual judgment.
# 4. For time-specific queries, prioritize temporal information in the content before considering timestamps.

# ---Query to be answered---
# {question}

# ---Conversation History---
# {history}

# ---Response Rules---

# - **Format**: Use multiple paragraphs with markdown formatting and section headings.
# - **Language**: Always respond in English, regardless of the language used in the question.
# - **Emphasis**: Highlight all referenced information using **bold text**.
# - **Continuity**: Maintain logical flow with the conversation history.

# - **Inline Citations**:
#     - Cite the source **immediately after** the referenced information using square brackets (e.g., [1], [2]).
#     - Every time a source is used, it must be cited—even if it has been cited earlier.

# - **Reference Mapping**:
#     - Assign each source file a unique reference number starting from 1, in the order of **first appearance** in the main answer.
#     - Maintain a mapping between source file names and their assigned reference numbers.
#     - Use this mapping consistently throughout the answer.

# - **References Section**:
#     - **Only include a "References" section if sources were actually cited in the main answer.**
#     - If no information was found or no sources were cited, do not include a "References" section.
#     - When including references, list only the **file names** (not full paths) that were cited in the main answer.
#     - List them in ascending order of their citation number.
#     - Each file name should appear **only once**.
#     - Do not include any file name that was not cited in the main answer.
#     - Ensure that every reference number used in the main answer appears exactly once in the "References" section.

# - **Integrity**:
#     - If the answer is unknown, say so clearly.
#     - Do not fabricate information or include anything not present in the Knowledge Base or conversation history."""

        print(
            "Domain:", domain, 
            "kb_name:", kb_name, 
            "knowledge_bases:", knowledge_bases,
            "question:", question, 
            "user_prompt:", user_prompt[:10], 
            "history:", history[:10]
        )
        if knowledge_bases:
            llm_summarize = AzureCustomLLM(temperature=0)
            results = {}
            kb_graph_refs = []

            # Initialize RAGs sequentially to avoid os.environ['NEO4J_DATABASE'] race condition
            # during parallel initialization. Query execution remains parallel below.
            rag_map = {}
            for kb in knowledge_bases:
                rag_map[kb] = await initialize_rag(domain=domain, kb_name=kb_name+kb)

            async def query_single_kb(kb):
                try:
                    rag = rag_map[kb]
                    response = await rag.aquery(
                        question if question else "",
                        param=QueryParam(
                            mode=mode,
                            top_k=2,
                            conversation_history=history,
                            user_prompt=user_prompt,
                            stream=False
                        )
                    )
                    return (kb, response, None)
                except Exception as e:
                    return (kb, None, str(e))

            tasks = [query_single_kb(kb) for kb in knowledge_bases]
            task_results = await asyncio.gather(*tasks)
            for kb, response, error in task_results:
                if error:
                    results[kb] = {"error": error}
                else:
                    results[kb] = response
                    kb_graph_refs.append(f"Knowledge Base: {kb}")

            # Summarize the results using AzureCustomLLM
            summary_prompt = user_prompt
            
            for kb, resp in results.items():
                summary_prompt += f"### Knowledge Base: {kb}\n"
                if isinstance(resp, dict) and 'error' in resp:
                    summary_prompt += f"Error: {resp['error']}\n"
                else:
                    summary_prompt += f"Response: {str(resp)}\n"
            summary_prompt += "\n---\nReferences:\n"
            for i, kb in enumerate(results.keys(), 1):
                summary_prompt += f"[{i}] {kb}\n"
            summary = llm_summarize._call(
                input=summary_prompt
            )

            
            
            # Parse references from the summary and generate download URLs
            original_kb_name = kb_name.split('/')[0] if '/' in kb_name else kb_name
            sources = []
            parsed_refs = parse_references_from_response(summary)  # Parse from summary, not response

            print(f"Parsed {len(parsed_refs)} references from LightRAG response")

            for ref in parsed_refs:
                citation = ref['citation_number']
                file_path = ref['file_path']
                
                print(f"Processing reference {citation}: {file_path}")
                
                download_url = generate_download_url_for_file(domain, original_kb_name, file_path,workspace_id=workspace_id,
                role_id=role_id)
                
                if download_url:
                    sources.append({
                        "file_name": f"{citation} {os.path.basename(file_path)}",
                        "download_url": download_url
                    })
                    print(f"Added source: {citation} {os.path.basename(file_path)}")
                else:
                    print(f"Skipped reference {citation}: Not a valid file or file not found")

            print(f"Final sources count: {len(sources)}")

            # Remove the References section from summary before sending to UI
            #summary_clean = re.sub(r'\n*###\s*References\s*\n.*', '', summary, flags=re.DOTALL).strip()
            summary_clean = re.sub(r'(?i)\n*##+\s*References[\s\S]*$', '', summary).strip()
            print(f" Removed References section. Clean response length: {len(summary_clean)}")

            return {
                "LightRAG": summary_clean,
                "response": summary_clean,
                "sources": sources
            }
        else:
            rag = await initialize_rag(domain=domain, kb_name=kb_name)
            response = await rag.aquery(
                question if question else "",
                param=QueryParam(
                    mode=mode, 
                    top_k=2, 
                    conversation_history=history, 
                    user_prompt=user_prompt, 
                    stream=False
                )
            )

            

            
            
            # Parse references from the response and generate download URLs
            sources = []
            
            # DEBUG: Check response for References section
            print("="*80)
            print("🔍 DEBUG: Response Analysis")
            print(f"Response type: {type(response)}")
            print(f"Response length: {len(response) if isinstance(response, str) else 'N/A'}")
            print(f"Has ### References: {'### References' in response if isinstance(response, str) else False}")
            if isinstance(response, str) and '### References' in response:
                ref_section = response.split('### References')[1][:300]
                print(f"📄 References section found:\n{ref_section}")
            print("="*80)
            
            parsed_refs = parse_references_from_response(response)
            print(f"✓ Parsed {len(parsed_refs)} references")
            if parsed_refs:
                for ref in parsed_refs:
                    print(f"  - {ref}")
            
            # Extract original kb_name without workspace suffix
            original_kb_name = kb_name.split('/')[0] if '/' in kb_name else kb_name
            for ref in parsed_refs:
                citation = ref['citation_number']  # e.g., "[1]"
                file_path = ref['file_path']
                
                print(f"🔗 Processing {citation}: {file_path}")
                print(f"   Domain: {domain}, KB: {original_kb_name}, WorkspaceID: {workspace_id}, RoleID: {role_id}")
                
                # Use ORIGINAL kb_name for blob path (not the modified one)
                download_url = generate_download_url_for_file(domain, original_kb_name, file_path,workspace_id=workspace_id,
                role_id=role_id)
                
                if download_url:
                    sources.append({
                        "file_name": f"{citation} {os.path.basename(file_path)}",
                        "download_url": download_url
                    })
                    print(f"✓ Generated URL for {citation}")
                else:
                    print(f"❌ FAILED to generate URL for {citation} - file not found in blob storage")
            
            print(f" Final sources count: {len(sources)}")
            print("="*80)
            
            # Remove the References section from response before sending to UI
           # response_clean = re.sub(r'\n*###\s*References\s*\n.*', '', response, flags=re.DOTALL).strip()
            response_clean = re.sub(r'(?i)\n*##+\s*References[\s\S]*$', '', response).strip()
            print(f" Removed References section. Clean response length: {len(response_clean)}")
            
            # Return structured response similar to SSRS
            return {
                "LightRAG": response_clean,
                "response": response_clean,
                "sources": sources
            }
    except Exception as e:
            return {"error": str(e)}


    
async def index_file(ctx: Context, container_client, domain, kb_name, file_path):
    print(f"Indexing file started: {file_path}")
    try:
        blob_client = container_client.get_blob_client(file_path)
        blob_data = blob_client.download_blob().readall()
        # Pass Data to your indexing function here
        resp_rag = await lightrag_indexing_tool_new(Context, blob_data, domain, kb_name, file_path)
        print(f"Response from RAG Indexing file completed: {resp_rag}")
    except Exception as e:
        return f"error: {e}"
 
async def index_sub_industry(ctx: Context, container_name, domain, kb_name, task_id=None):
    try:
        industry = domain
        sub_industry = kb_name
        max_concurrent = 3
        file_status = []
        prefix = f"{industry}/{sub_industry}/"
 
        connection_string = os.getenv('AZURE_BLOB_STORAGE_CONNECTION_STRING')
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        container_client = blob_service_client.get_container_client(container_name or "")
       
 
        # List all files in the subfolder
        blob_list = container_client.list_blobs(name_starts_with=prefix)
        semaphore = asyncio.Semaphore(max_concurrent)
        async def sem_index_file(file_path):
            async with semaphore:
                await index_file(Context, container_client, industry, sub_industry, file_path)
       
        tasks = [sem_index_file(blob.name) for blob in blob_list]
        await asyncio.gather(*tasks)
    except Exception as e:
        return f"Error indexing {industry}/{sub_industry}: {e}"
   
@mcp.tool()  
def start_workspace_indexing(ctx: Context, container_name: Optional[str] = None, domain: Optional[str] = None, kb_name: Optional[List[str]] = None, file_path: Optional[str] = None):
    try:
        industry = domain
        sub_industries = kb_name
        for sub_industry in sub_industries:
            task_id = f"{industry}_{sub_industry}_{int(time.time())}"
            print(f"Starting indexing for {industry}/{sub_industry} with task_id {task_id}")
            task = asyncio.create_task(index_sub_industry(Context, container_name, industry, sub_industry, task_id))
        return {"status": "success"}
    except Exception as e:
        return {"error": str(e)}
   
 
@mcp.tool()
async def index_uploaded_files(ctx: Context, container_name, domain, kb_names, file_names):
    """
    For each kb_name and file_name, generate path domain/kb_name/file_name and index it.
    """
    try:
        if not file_names:
            return {"error": "Files are required for indexing."}
       
        connection_string = os.getenv('AZURE_BLOB_STORAGE_CONNECTION_STRING')
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        container_client = blob_service_client.get_container_client(container_name or "")
 
        tasks = []
        for kb in kb_names:
            for file_name in file_names:
                file_path = f"{domain}/{kb}/{file_name}"
                print(f"Scheduling indexing for: {file_path}")
                tasks.append(index_file(ctx, container_client, domain, kb, file_path))
       
        results = await asyncio.gather(*tasks)
        return {"status": "success", "tasks": results}
    except Exception as e:
        return {"error": str(e)}
   
async def lightrag_indexing_tool_new(ctx: Context, blob_data, domain, kb_name, file_path) -> dict:
    """
    Index all .txt files in the specified domain/KB directory.
    """
    try:
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        ext = file_path.lower().split('.')[-1]
        content = None
       
        if ext == 'txt':
            content = blob_data.decode('utf-8', errors='ignore')
        # elif ext == 'pdf':
        #     import fitz  # PyMuPDF
        #     from paddleocr import PaddleOCR
        #     from PIL import Image
        #     import numpy as np
        #     ocr = PaddleOCR(use_angle_cls=True, lang='en')
        #     content = ""
        #     with fitz.open(stream=blob_data, filetype="pdf") as doc:
        #         all_text = []
        #         for page_num in range(len(doc)):
        #             page = doc.load_page(page_num)
        #             # Use get_pixmap for PyMuPDF >= 1.18, fallback to getPixmap for older
        #             if hasattr(page, 'get_pixmap'):
        #                 pix = page.get_pixmap()
        #             elif hasattr(page, 'getPixmap'):
        #                 pix = page.getPixmap()
        #             else:
        #                 raise RuntimeError('No get_pixmap or getPixmap method found on PyMuPDF Page object')
        #             img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        #             img_np = np.array(img)
        #             result = ocr.predict(img_np)
        #             page_text = " ".join(result[0]['rec_texts'])
        #             print(f"OCR result for {file_path} page {page_num}:", result[0]['rec_texts'])
        #             all_text.append(page_text)
        #         content = "\n".join(all_text)
        #     print(content)
        elif ext in ['docx']:
            # Ensure blob data represents raw DOCX/DOC bytes (not base64 text)
            def _ensure_doc_bytes(data: bytes | str) -> bytes:
                try:
                    if isinstance(data, (bytes, bytearray)):
                        sample = bytes(data[:32])
                        # If the start looks ASCII, try strict base64 decode
                        is_ascii_like = all(32 <= b <= 126 or b in (9, 10, 13) for b in sample)
                        if is_ascii_like:
                            try:
                                decoded = base64.b64decode(data, validate=True)
                                # DOCX files are ZIPs starting with PK\x03\x04
                                if decoded.startswith(b"PK\x03\x04"):
                                    return decoded
                            except Exception:
                                pass
                        return bytes(data)
                    elif isinstance(data, str):
                        # Handle data URLs or raw base64 strings
                        if data.startswith("data:"):
                            try:
                                b64_part = data.split("base64,", 1)[-1]
                                return base64.b64decode(b64_part, validate=True)
                            except Exception:
                                return data.encode("utf-8")
                        try:
                            return base64.b64decode(data, validate=True)
                        except Exception:
                            return data.encode("utf-8")
                    else:
                        return bytes(data)
                except Exception:
                    return bytes(data) if isinstance(data, (bytes, bytearray)) else str(data).encode("utf-8")

            doc_bytes = _ensure_doc_bytes(blob_data)
            doc = Document(io.BytesIO(doc_bytes))
            content = "\n".join([p.text for p in doc.paragraphs])
        elif ext == 'doc':
            # Attempt Windows COM-based conversion (.doc -> .docx) if Word is available
            def _convert_doc_to_docx_bytes(doc_bytes: bytes) -> bytes | None:
                try:
                    import win32com.client as win32
                    import pythoncom
                    # Write input .doc to a temp file
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as f_in:
                        f_in.write(doc_bytes)
                        in_path = f_in.name
                    out_path = in_path[:-4] + '.docx'
                    try:
                        pythoncom.CoInitialize()
                        word = win32.Dispatch('Word.Application')
                        word.Visible = False
                        doc = word.Documents.Open(in_path)
                        wdFormatXMLDocument = 12
                        doc.SaveAs(out_path, FileFormat=wdFormatXMLDocument)
                        doc.Close(False)
                        word.Quit()
                        pythoncom.CoUninitialize()
                        with open(out_path, 'rb') as f_out:
                            return f_out.read()
                    finally:
                        # Cleanup temp files
                        try:
                            import os
                            if os.path.exists(in_path):
                                os.remove(in_path)
                            if os.path.exists(out_path):
                                os.remove(out_path)
                        except Exception:
                            pass
                except Exception:
                    return None

            # Detect typical .doc OLE header
            hdr = bytes(blob_data[:8]) if isinstance(blob_data, (bytes, bytearray)) else b''
            if hdr.startswith(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1") or hdr.startswith(b"0M8R4KG"):
                converted = _convert_doc_to_docx_bytes(blob_data if isinstance(blob_data, (bytes, bytearray)) else blob_data.encode('utf-8'))
                if converted:
                    doc = Document(io.BytesIO(converted))
                    content = "\n".join([p.text for p in doc.paragraphs])
                else:
                    # Fallback: extract readable ASCII sequences from .doc bytes
                    try:
                        import re
                        raw = blob_data if isinstance(blob_data, (bytes, bytearray)) else str(blob_data).encode('utf-8', errors='ignore')
                        # Replace non-printable with spaces, keep basic punctuation
                        text = raw.decode('latin-1', errors='ignore')
                        # Find sequences of 5+ printable chars
                        blocks = re.findall(r"[\x20-\x7E]{5,}", text)
                        content = "\n".join(blocks)
                        if not content or len(content.strip()) < 20:
                            return {"error": "Legacy .doc detected and conversion unavailable. Could not extract sufficient text. Please convert to .docx and re-upload."}
                    except Exception:
                        return {"error": "Legacy .doc detected and conversion unavailable. Could not extract text reliably. Please convert to .docx and re-upload."}
            else:
                return {"error": "Unknown .doc content format. Only .docx is supported for indexing."}
        else:
            return {"error": f"Unsupported file type: {ext}"}
        def chunk_text(text, chunk_size=2000):
            # Simple chunking by character count
            return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
 
        chunks = chunk_text(content, chunk_size=2000)
        for idx, chunk in enumerate(chunks):
            await rag.ainsert(input=chunk, file_paths=[file_path])
            await ctx.debug(f"Progress: {idx+1}/{len(chunks)}")
        return {"status": "success", "file": file_path, "chunks": len(chunks)}
    except Exception as e:
        return {"error": str(e)}
   
def uploaded_by_username(uploaded_by_value: str | None) -> str | None:
    """
    If uploaded_by is a numeric user_id, resolve it to 'First Last' from users table.
    Otherwise return as-is.
    """
    try:
        if uploaded_by_value is None:
            return None
        if str(uploaded_by_value).isdigit():
            conn = psycopg2.connect(
                host=os.environ["POSTGRES_HOST"],
                user=os.environ["POSTGRES_USER"],
                password=os.environ["POSTGRES_PASSWORD"],
                dbname=os.environ["POSTGRESQL_DATABASE_DATABASE_2"]
            )
            cur = conn.cursor()
            cur.execute(
                "SELECT COALESCE(first_name,'') AS fn, COALESCE(last_name,'') AS ln FROM users WHERE user_id = %s",
                (int(uploaded_by_value),)
            )
            row = cur.fetchone()
            cur.close()
            conn.close()
            if row:
                fn, ln = row
                full_name = (fn + " " + ln).strip()
                return full_name if full_name else str(uploaded_by_value)
        return uploaded_by_value
    except Exception as e:
        print(f"Warning: could not resolve uploaded_by to name: {e}")
        return uploaded_by_value

def create_file_task_record(container_name, upload_path, domain, kb_name, file_path, workspace_id, status="uploading", file_size=None, uploaded_by=None):
    try:
        # Resolve uploaded_by (user_id -> full name)
        uploaded_by_resolved = uploaded_by_username(uploaded_by)

        conn = psycopg2.connect(
            host=os.environ["POSTGRES_HOST"],
            user=os.environ["POSTGRES_USER"],
            password=os.environ["POSTGRES_PASSWORD"],
            dbname=os.environ["POSTGRESQL_DATABASE_DATABASE_2"]
        )
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO file_tasks
            (container_name, upload_path, domain, kb_name, file_path, workspace_id, status, file_size, uploaded_by, created_at, updated_at)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, NOW(), NOW())
            RETURNING id;
        """, (container_name, upload_path, domain, kb_name, file_path, workspace_id, status, file_size, uploaded_by_resolved))
        task_id = cur.fetchone()[0]
        conn.commit()
        cur.close()
        conn.close()
        return task_id
    except Exception as e:
        print(f"Error creating file_tasks record: {e}")
        return None
 
def update_file_task_status(task_id, status):
    try:
        conn = psycopg2.connect(
            host=os.environ["POSTGRES_HOST"],
            user=os.environ["POSTGRES_USER"],
            password=os.environ["POSTGRES_PASSWORD"],
            dbname=os.environ["POSTGRESQL_DATABASE_DATABASE_2"]
        )
        cur = conn.cursor()
        cur.execute("""
            UPDATE file_tasks SET status=%s, updated_at=NOW() WHERE id=%s;
        """, (status, task_id))
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"Error updating file_tasks status: {e}")

def _estimate_content_size_bytes(fcontent) -> int | None:
    """
    Best-effort byte-size estimation for incoming file content.
    """
    try:
        if isinstance(fcontent, (bytes, bytearray)):
            return len(fcontent)
        if isinstance(fcontent, str):
            # try base64 first
            try:
                return len(base64.b64decode(fcontent, validate=True))
            except Exception:
                # fallback to utf-8 encoding size
                return len(fcontent.encode("utf-8"))
        # last resort
        return len(bytes(fcontent))
    except Exception:
        return None   
       

def _format_size_with_unit(size_bytes: int | None) -> str | None:
    """
    Convert a byte size into a human-readable string with units.
    Returns values with up to two decimal places using binary units (KB, MB, GB).

    Examples:
    - 1536 -> "1.50 KB"
    - 1048576 -> "1.00 MB"
    - 1073741824 -> "1.00 GB"
    """
    try:
        if size_bytes is None:
            return None
        # Use binary measurement (KiB, MiB, GiB) but display as KB/MB/GB
        KB = 1024
        MB = KB * 1024
        GB = MB * 1024

        if size_bytes < KB:
            # Keep bytes as-is; although request focuses on KB/MB/GB,
            # small files will still be accurately represented.
            return f"{size_bytes} Bytes"
        elif size_bytes < MB:
            return f"{size_bytes / KB:.2f} KB"
        elif size_bytes < GB:
            return f"{size_bytes / MB:.2f} MB"
        else:
            return f"{size_bytes / GB:.2f} GB"
    except Exception:
        # Fallback to raw bytes string if something goes wrong
        return str(size_bytes) if size_bytes is not None else None
 
# Reusable upload function (keeps same functionality as your original upload_files_and_get_urls)
async def upload_files_and_get_urls(container_name: str, file_path: str, file_names: list, file_contents: list, expiry_years: int = 10):
   
    print("File names to upload:", file_names)
    print("File contents types:", [type(fc) for fc in file_contents])
    print("File contents sizes:", [len(fc) if isinstance(fc, (bytes, str)) else 'N/A' for fc in file_contents])
    """
    Uploads files to Azure Blob Storage and returns their long-lived SAS download URLs.
 
    Args:
        container_name (str): Name of the container where file will be uploaded.
        file_path (str): Path of the uploaded file.
        file_names (list): List of file names to upload.
        file_contents (list): List of file contents (bytes or str).
        expiry_years (int): Years until the SAS token expires (default: 10).
 
    Returns:
        dict: {file_name: download_url or error}
    """
    connection_string = os.getenv('AZURE_BLOB_STORAGE_CONNECTION_STRING')
 
    if not connection_string or not container_name:
        return {"error": "Azure Blob Storage configuration is missing."}
 
    if not (len(file_names) == len(file_contents)):
        return {"error": "file_names and file_contents must have the same length."}
 
    try:
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        container_client = blob_service_client.get_container_client(container_name)
        result = {}
        account_name = str(blob_service_client.account_name) if blob_service_client.account_name else ""
        if not account_name:
            return {"error": "Could not determine Azure Storage account name."}
        for fname, fcontent in zip(file_names, file_contents):
            # Normalize input content to raw bytes
            original_type = type(fcontent)
 
            # Handle data URL strings
            if isinstance(fcontent, str):
                if fcontent.startswith("data:"):
                    try:
                        _, b64data = fcontent.split(",", 1)
                        fcontent = b64data
                    except ValueError:
                        pass
                # Try base64 decode; if that fails, treat as UTF-8 text bytes
                try:
                    fcontent = base64.b64decode(fcontent, validate=True)
                except Exception:
                    fcontent = fcontent.encode("utf-8")
 
            # Detect and decode base64 provided as bytes
            elif isinstance(fcontent, (bytes, bytearray)):
                sample = bytes(fcontent[:10])
                is_ascii_like = all(32 <= b <= 126 or b in (9, 10, 13) for b in sample)
                if is_ascii_like or sample.startswith(b"JVBER"):
                    try:
                        decoded = base64.b64decode(fcontent, validate=True)
                        if decoded.startswith(b"%PDF"):
                            fcontent = decoded
                    except Exception:
                        pass
            else:
                # Fallback conversion
                try:
                    fcontent = bytes(fcontent)
                except Exception:
                    return {"error": f"Unsupported content type for {fname}: {original_type}"}
 
            # Choose content type based on file extension
            ext = os.path.splitext(fname)[1].lower()

            # For text-like files, decode base64 payloads that arrive as ASCII bytes.
            if ext in (".txt", ".log", ".csv", ".md", ".json") and isinstance(fcontent, (bytes, bytearray)):
                try:
                    sample = bytes(fcontent[:64])
                    is_ascii_like = all(32 <= b <= 126 or b in (9, 10, 13) for b in sample)
                    if is_ascii_like:
                        decoded = base64.b64decode(fcontent, validate=True)
                        if decoded:
                            fcontent = decoded
                except Exception:
                    # Keep original bytes when payload is not base64-encoded.
                    pass

            # For DOC/DOCX specifically, ensure we upload raw binary bytes (not base64 text)
            if ext in (".doc", ".docx"):
                try:
                    if isinstance(fcontent, str):
                        if fcontent.startswith("data:"):
                            b64_part = fcontent.split("base64,", 1)[-1]
                            fcontent = base64.b64decode(b64_part, validate=True)
                        else:
                            try:
                                fcontent = base64.b64decode(fcontent, validate=True)
                            except Exception:
                                fcontent = fcontent.encode("utf-8")
                    elif isinstance(fcontent, (bytes, bytearray)):
                        # If bytes look like ASCII base64, try decoding generically
                        sample = bytes(fcontent[:12])
                        is_ascii_like = all(32 <= b <= 126 or b in (9, 10, 13) for b in sample)
                        if is_ascii_like:
                            try:
                                decoded = base64.b64decode(fcontent, validate=False)
                                # Accept decoded if it matches DOCX ZIP (PK..) or DOC OLE (D0 CF 11 E0 ...)
                                if (
                                    decoded.startswith(b"PK\x03\x04") or
                                    decoded.startswith(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1")
                                ):
                                    fcontent = decoded
                            except Exception:
                                # fall through, keep original bytes
                                pass
                        # Ensure type is bytes
                        fcontent = bytes(fcontent)
                    else:
                        fcontent = bytes(fcontent)
                except Exception:
                    fcontent = bytes(fcontent) if isinstance(fcontent, (bytes, bytearray)) else str(fcontent).encode("utf-8")

            if ext == ".pdf":
                content_settings = ContentSettings(content_type="application/pdf")
            elif ext in (".txt", ".log"):
                content_settings = ContentSettings(content_type="text/plain")
            elif ext == ".docx":
                content_settings = ContentSettings(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            elif ext == ".doc":
                content_settings = ContentSettings(content_type="application/msword")
            else:
                content_settings = ContentSettings(content_type="application/octet-stream")
 
            # Upload
            blob_path = file_path + f"/{fname}"
            blob_client = container_client.get_blob_client(blob_path)
            try:
                blob_client.upload_blob(fcontent, overwrite=True, content_settings=content_settings)
                expiry = datetime.now() + timedelta(days=365 * expiry_years)
                sas_token = generate_blob_sas(
                    account_name=account_name,
                    container_name=container_name,
                    blob_name=blob_path,
                    account_key=blob_service_client.credential.account_key,
                    permission=BlobSasPermissions(read=True),
                    expiry=expiry,
                    # Force download in browser
                    content_disposition=f'attachment; filename="{os.path.basename(blob_path)}"'
                )
                download_url = f"https://{account_name}.blob.core.windows.net/{container_name}/{blob_path}?{sas_token}"
                print(f"Uploaded {fname} to {download_url}")
                result[fname] = download_url
                print(f"Upload successful for {fname}")
            except Exception as e:
                result[fname] = f"Error: {str(e)}, Fcontent: {type(fcontent)}"
        return result
    except Exception as e:
        return {"error": str(e)}
   
 
async def lightrag_indexing_tool(
    container_name: Optional[str] = None, 
    domain: Optional[str] = None, 
    kb_name: Optional[str] = None, 
    file_path: Optional[str] = None
) -> dict:
    """
    Index all .txt files in the specified domain/KB directory.
    """
    try:
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        if not file_path:
            return {"error": "file_path is required"}
        ext = file_path.lower().split('.')[-1]
        content = None
        # Read file from Azure Blob Storage
        connection_string = os.getenv('AZURE_BLOB_STORAGE_CONNECTION_STRING')
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        print(file_path)
        container_client = blob_service_client.get_container_client(container_name or "")
        blob_client = container_client.get_blob_client(file_path)
        blob_data = blob_client.download_blob().readall()
        print(blob_data[:100])
        if ext == 'txt':
            content = blob_data.decode('utf-8', errors='ignore')
        elif ext == 'pdf':
            # Set up credentials and client
            try:
                endpoint = os.getenv('AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT')
                api_key = os.getenv('AZURE_DOCUMENT_INTELLIGENCE_KEY')
                if not endpoint or not api_key:
                    return {"error": "Azure Document Intelligence endpoint or key not set in environment variables."}
                doc_client = DocumentIntelligenceClient(endpoint, AzureKeyCredential(api_key))
 
                # Pass the stream directly to Document Intelligence
                poller = doc_client.begin_analyze_document(
                    "prebuilt-read",
                    body=AnalyzeDocumentRequest(bytes_source=blob_data),
                    locale="en-US"
                )
                result = poller.result()
 
                content = result.content
            except Exception as e:
                return {"error": f"Failed to process PDF with Document Intelligence: {e}"}
        elif ext in ['docx']:
            # Ensure blob data represents raw DOCX/DOC bytes (not base64 text)
            def _ensure_doc_bytes(data: bytes | str) -> bytes:
                try:
                    if isinstance(data, (bytes, bytearray)):
                        sample = bytes(data[:32])
                        is_ascii_like = all(32 <= b <= 126 or b in (9, 10, 13) for b in sample)
                        if is_ascii_like:
                            try:
                                decoded = base64.b64decode(data, validate=True)
                                if decoded.startswith(b"PK\x03\x04"):
                                    return decoded
                            except Exception:
                                pass
                        return bytes(data)
                    elif isinstance(data, str):
                        if data.startswith("data:"):
                            try:
                                b64_part = data.split("base64,", 1)[-1]
                                return base64.b64decode(b64_part, validate=True)
                            except Exception:
                                return data.encode("utf-8")
                        try:
                            return base64.b64decode(data, validate=True)
                        except Exception:
                            return data.encode("utf-8")
                    else:
                        return bytes(data)
                except Exception:
                    return bytes(data) if isinstance(data, (bytes, bytearray)) else str(data).encode("utf-8")

            doc_bytes = _ensure_doc_bytes(blob_data)
            doc = Document(io.BytesIO(doc_bytes))
            content = "\n".join([p.text for p in doc.paragraphs])
        elif ext == 'doc':
            # Attempt Windows COM-based conversion (.doc -> .docx) if Word is available
            def _convert_doc_to_docx_bytes(doc_bytes: bytes) -> bytes | None:
                try:
                    import win32com.client as win32
                    import pythoncom
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as f_in:
                        f_in.write(doc_bytes)
                        in_path = f_in.name
                    out_path = in_path[:-4] + '.docx'
                    try:
                        pythoncom.CoInitialize()
                        word = win32.Dispatch('Word.Application')
                        word.Visible = False
                        doc = word.Documents.Open(in_path)
                        wdFormatXMLDocument = 12
                        doc.SaveAs(out_path, FileFormat=wdFormatXMLDocument)
                        doc.Close(False)
                        word.Quit()
                        pythoncom.CoUninitialize()
                        with open(out_path, 'rb') as f_out:
                            return f_out.read()
                    finally:
                        try:
                            import os
                            if os.path.exists(in_path):
                                os.remove(in_path)
                            if os.path.exists(out_path):
                                os.remove(out_path)
                        except Exception:
                            pass
                except Exception:
                    return None

            hdr = bytes(blob_data[:8]) if isinstance(blob_data, (bytes, bytearray)) else b''
            if hdr.startswith(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1") or hdr.startswith(b"0M8R4KG"):
                converted = _convert_doc_to_docx_bytes(blob_data if isinstance(blob_data, (bytes, bytearray)) else blob_data.encode('utf-8'))
                if converted:
                    doc = Document(io.BytesIO(converted))
                    content = "\n".join([p.text for p in doc.paragraphs])
                else:
                    # Fallback: try extracting readable ASCII sequences, so indexing doesn’t stop
                    try:
                        import re
                        raw = blob_data if isinstance(blob_data, (bytes, bytearray)) else str(blob_data).encode('utf-8', errors='ignore')
                        text = raw.decode('latin-1', errors='ignore')
                        blocks = re.findall(r"[\x20-\x7E]{5,}", text)
                        content = "\n".join(blocks)
                        if not content or len(content.strip()) < 20:
                            return {"error": "Legacy .doc detected and conversion unavailable. Could not extract sufficient text. Please convert to .docx and re-upload."}
                    except Exception:
                        return {"error": "Legacy .doc detected and conversion unavailable. Could not extract text reliably. Please convert to .docx and re-upload."}
            else:
                return {"error": "Unknown .doc content format. Only .docx is supported for indexing."}
        else:
            return {"error": f"Unsupported file type: {ext}"}
        def chunk_text(text, chunk_size=2000):
            # Simple chunking by character count
            return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
 
        chunks = chunk_text(content, chunk_size=2000)
        for idx, chunk in enumerate(chunks):
            await rag.ainsert(input=chunk, file_paths=[file_path])
            #await ctx.debug(f"Progress: {idx+1}/{len(chunks)}")
        return {"status": "success", "file": file_path, "chunks": len(chunks)}
    except Exception as e:
        return {"error": str(e)}    
 
# Combined MCP tool: upload + indexing (both background)
@mcp.tool()
async def upload_and_index_tool(
    ctx: Context,
    workspace_id: Optional[str] = None,
    container_name: Optional[str] = None,
    upload_path: Optional[str] = None,
    file_names: Optional[List[str]] = None,
    file_contents: Optional[List[bytes]] = None,
    domain: Optional[str] = None,
    kb_name: Optional[str] = None,
    user_id: Optional[str] = None,
    expiry_years: int = 10
) -> dict:
    """
    Orchestrates uploading files to Azure Blob Storage and then indexing them.
    - Creates a single task row (file_tasks) with status 'uploading'
    - Immediately returns to client with task_id (avoids MCP timeout)
    - Runs the upload in background. On upload success => status 'uploaded'
    - Then starts indexing in background (status 'indexing') using the same indexing logic you had
    - Updates status to 'indexed' or 'failed'
    """
    if user_id is None:
        return {"status": "error", "error": "user_id cannot be null"}

    # --- JWT-based authentication and workspace-user mapping check (copied from list_workspace_users) ---
    # Validate workspace_id presence
    # if not workspace_id:
    #     await ctx.debug("workspace_id is required for authentication.")
    #     return {"error": "workspace_id is required for authentication."}

    # # Validate user access to workspace
    # valid, err = validate_user_workspace_access(workspace_id=workspace_id)
    # if not valid:
    #     await ctx.debug(f"User not mapped to workspace: {err}")
    #     return {"error": err}

    # # Enforce JWT-based access: only allow if user is mapped to the workspace
    # request = request_var.get(None)
    # if not request or not hasattr(request.state, "jwt_claims"):
    #     await ctx.debug("Unauthorized: JWT claims not found in request context")
    #     return {"error": "Unauthorized: JWT claims not found in request context"}
    # claims = request.state.jwt_claims
    # jwt_user_id = claims.get("user_id") or claims.get("sub")
    # if not jwt_user_id:
    #     await ctx.debug("Unauthorized: user_id not found in token claims")
    #     return {"error": "Unauthorized: user_id not found in token claims"}

    jwt_user_id = user_id

    # Check if user is mapped to this workspace
    session = db.Session()
    try:
        user_map = session.query(db.UserMap).filter_by(workspace_id=workspace_id, user_id=jwt_user_id, is_active=True).first()
        if not user_map:
            session.close()
            await ctx.debug("You are not authorized to access this workspace.")
            return {"error": "You are not authorized to access this workspace."}
    except Exception as e:
        session.close()
        await ctx.debug(f"Error during workspace-user mapping check: {e}")
        return {"error": str(e)}
    finally:
        pass

    # Validate input lists for multi-file flow
    if not file_names or not file_contents or len(file_names) != len(file_contents):
        await ctx.debug("file_names and file_contents are required and must match in length.")
        return {
            "message": "Invalid input: file_names and file_contents must be provided with equal lengths.",
            "status": "error"
        }

    # Create one background job per file, each with its own task row following the same state machine
    tasks_summary = []

    async def background_upload_then_index_single(fname: str, fcontent: bytes, fpath: str, tid: Optional[int]):
        try:
            if not tid:
                print(f"No task id created for file {fname}; aborting background flow.")
                return

            # ---- BACKGROUND UPLOAD (per-file) ----
            update_file_task_status(tid, "uploading")
            upload_result = await upload_files_and_get_urls(
                container_name,
                upload_path or "",
                [fname],
                [fcontent],
                expiry_years=expiry_years,
            )

            # Inspect upload result for this specific file
            per_file_error = False
            if isinstance(upload_result, dict):
                if upload_result.get("error"):
                    per_file_error = True
                else:
                    v = upload_result.get(fname)
                    if isinstance(v, str) and v.startswith("Error:"):
                        per_file_error = True

            if per_file_error:
                update_file_task_status(tid, "failed")
                print("Upload failed for task_id:", tid, "result:", upload_result)
                return

            # Upload success
            update_file_task_status(tid, "uploaded")
            print("Upload complete for task_id:", tid)

            # ---- START INDEXING (per-file) ----
            update_file_task_status(tid, "indexing")

            try:
                if not fpath:
                    update_file_task_status(tid, "failed")
                    return

                result = await lightrag_indexing_tool(
                    container_name=container_name,
                    domain=domain,
                    kb_name=kb_name,
                    file_path=fpath,
                )

                # Handle responses consistently: treat certain benign errors as indexed
                if isinstance(result, dict) and result.get("error"):
                    error_msg = str(result.get("error"))
                    if (
                        'already exists' in error_msg
                        or 'No new unique documents' in error_msg
                        or 'No documents to process' in error_msg
                    ):
                        print("File indexing already completed or nothing to index for task_id:", tid)
                        update_file_task_status(tid, "indexed")
                    else:
                        print("Indexing failed for task_id:", tid, "error:", error_msg)
                        update_file_task_status(tid, "failed")
                        return
                else:
                    update_file_task_status(tid, "indexed")
                    print("Indexing complete for task_id:", tid)

            except Exception as e:
                # handle exceptions from indexing
                tb = traceback.format_exc()
                error_msg = str(e) or "Unknown error (exception has no message)"
                if 'already exists' in error_msg or 'No new unique documents' in error_msg or 'No documents to process' in error_msg:
                    print("File indexing already completed or nothing to index for task_id:", tid)
                    update_file_task_status(tid, "indexed")
                else:
                    print(f"Error during background indexing: {error_msg}\nTraceback:\n{tb}")
                    update_file_task_status(tid, "failed")
                    return

        except Exception as e:
            # Any top-level unexpected exception in the background flow
            tb = traceback.format_exc()
            print(f"Unexpected background error for task_id {tid}: {e}\n{tb}")
            if tid:
                update_file_task_status(tid, "failed")

    # Kick off one background coroutine per file
    print("Starting background upload and indexing tasks for files:", file_names)
    print("length file_contents", len(file_contents))
    for fname, fcontent in zip(file_names, file_contents):
        per_file_path = f"{upload_path}/{fname}" if upload_path and fname else None
        # Compute human-readable size with units for storage in file_tasks.file_size
        _bytes = _estimate_content_size_bytes(fcontent)
        estimated_size = _format_size_with_unit(_bytes)
        tid = create_file_task_record(
            container_name,
            upload_path,
            domain,
            kb_name,
            per_file_path,
            workspace_id,
            status="uploading",
            file_size=estimated_size,
            uploaded_by=user_id,  # use user_id as "uploaded_by"
        )
        tasks_summary.append({
            "file_name": fname,
            "file_path": per_file_path,
            "task_id": tid,
        })
        asyncio.create_task(background_upload_then_index_single(fname, fcontent, per_file_path, tid))

    # Immediately inform client that tasks started (avoid MCP timeout)
    await ctx.debug("Upload(s) started in background. Use the returned task_ids to poll status.")
    return {
        "message": "Upload and indexing started in background for all files.",
        "status": "background",
        "tasks": tasks_summary,
    }
 
 
# MCP tool to check indexing status
@mcp.tool()
async def check_specific_indexing_status(
    ctx: Context,
    task_ids: Optional[list] = None
) -> dict:
    """
    Check the status of one or more indexing tasks by task_ids (list of int or str).
    Returns a list of status dicts for each task_id.
    """
    if not task_ids:
        return {"error": "task_ids is required (list of task ids)"}
    # Allow backward compatibility: if a single int is passed, treat as list
    if isinstance(task_ids, (int, str)):
        task_ids = [task_ids]
    try:
        conn = psycopg2.connect(
            host=os.environ["POSTGRES_HOST"],
            user=os.environ["POSTGRES_USER"],
            password=os.environ["POSTGRES_PASSWORD"],
            dbname=os.environ["POSTGRESQL_DATABASE_DATABASE_2"]
        )
        cur = conn.cursor()
        # Prepare query for multiple ids
        format_strings = ','.join(['%s'] * len(task_ids))
        cur.execute(f"""
            SELECT id, status, created_at, updated_at, file_path, domain, kb_name, workspace_id, file_size, uploaded_by FROM file_tasks
            WHERE id IN ({format_strings})
        """, tuple(task_ids))
        rows = cur.fetchall()
        cur.close()
        conn.close()
        found = {row[0]: row for row in rows}
        results = []
        for tid in task_ids:
            row = found.get(int(tid)) if isinstance(tid, (int, str)) and str(tid).isdigit() else None
            if row:
                task_id, status, created_at, updated_at, file_path, domain, kb_name, workspace_id, file_size, uploaded_by = row
                results.append({
                    "task_id": task_id,
                    "status": status,
                    "created_at": str(created_at),
                    "updated_at": str(updated_at),
                    "file": file_path,
                    "domain": domain,
                    "kb_name": kb_name,
                    "workspace_id": workspace_id,
                    "file_size": file_size,
                    "uploaded_by": uploaded_by,
                })
            else:
                results.append({
                    "message": "No indexing task found for the specified task_id.",
                    "task_id": tid
                })
        return {"results": results}
    except Exception as e:
        return {"error": str(e)}
    

@mcp.tool()
async def check_indexing_status_by_workspace(
    ctx: Context,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
    limit: int = 100,
    include_counts: bool = True
) -> dict:
    """
    List indexing tasks for a given workspace_id.
    Returns a list of tasks (most recent first) and optional status counts.
    """
    if user_id is None:
        return {"status": "error", "error": "user_id cannot be null"}

    # --- JWT-based authentication and workspace-user mapping check (copied from list_workspace_users) ---
    if not workspace_id:
        return {"error": "workspace_id is required"}

    # Validate user access to workspace
    valid, err = validate_user_workspace_access(workspace_id=workspace_id)
    if not valid:
        await ctx.debug(f"User not mapped to workspace: {err}")
        return {"error": err}

    # Enforce JWT-based access: only allow if user is mapped to the workspace
    request = request_var.get(None)
    if not request or not hasattr(request.state, "jwt_claims"):
        await ctx.debug("Unauthorized: JWT claims not found in request context")
        return {"error": "Unauthorized: JWT claims not found in request context"}
    claims = request.state.jwt_claims
    jwt_user_id = claims.get("user_id") or claims.get("sub")
    if not jwt_user_id:
        await ctx.debug("Unauthorized: user_id not found in token claims")
        return {"error": "Unauthorized: user_id not found in token claims"}

    # Check if user is mapped to this workspace
    session = db.Session()
    try:
        user_map = session.query(db.UserMap).filter_by(workspace_id=workspace_id, user_id=jwt_user_id, is_active=True).first()
        if not user_map:
            session.close()
            await ctx.debug("You are not authorized to access this workspace.")
            return {"error": "You are not authorized to access this workspace."}
    except Exception as e:
        session.close()
        await ctx.debug(f"Error during workspace-user mapping check: {e}")
        return {"error": str(e)}
    finally:
        pass

    try:
        conn = psycopg2.connect(
            host=os.environ["POSTGRES_HOST"],
            user=os.environ["POSTGRES_USER"],
            password=os.environ["POSTGRES_PASSWORD"],
            dbname=os.environ["POSTGRESQL_DATABASE_DATABASE_2"]
        )
        cur = conn.cursor()

        # Optional filter by uploaded_by derived from provided user_id
        filter_uploaded_by = None
        if user_id:
            # Resolve numeric user_id to the stored uploaded_by (full name) if possible
            filter_uploaded_by = uploaded_by_username(str(user_id))

        # Build query with optional filter
        base_query = """
            SELECT id, status, created_at, updated_at, file_path, domain, kb_name, file_size, uploaded_by, container_name
            FROM file_tasks
            WHERE workspace_id = %s
        """
        params = [workspace_id]
        if filter_uploaded_by is not None:
            base_query += " AND uploaded_by = %s"
            params.append(filter_uploaded_by)

        base_query += " ORDER BY created_at DESC LIMIT %s"
        params.append(limit)

        # Fetch recent tasks for this workspace (optionally filtered by user)
        cur.execute(base_query, tuple(params))
        rows = cur.fetchall()

        # Best-effort: prepare Azure SAS prerequisites
        connection_string = os.getenv('AZURE_BLOB_STORAGE_CONNECTION_STRING')
        account_name = None
        account_key = None
        endpoint_suffix = 'core.windows.net'
        if connection_string:
            try:
                _bsc = BlobServiceClient.from_connection_string(connection_string)
                account_name = str(_bsc.account_name) if _bsc.account_name else None
            except Exception:
                account_name = None

            def _conn_val(cs: str, key: str):
                try:
                    parts = [p for p in cs.split(';') if p]
                    for p in parts:
                        if p.startswith(key + '='):
                            return p.split('=', 1)[1]
                except Exception:
                    return None
                return None

            if not account_name:
                account_name = _conn_val(connection_string, 'AccountName')
            endpoint_suffix = _conn_val(connection_string, 'EndpointSuffix') or endpoint_suffix
            account_key = _conn_val(connection_string, 'AccountKey')

        tasks = []
        for r in rows:
            # Unpack matching 10 selected columns (includes container_name)
            task_id, status, created_at, updated_at, file_path, domain, kb_name, file_size, uploaded_by, container_name = r

            # Generate per-file SAS URL if possible (non-fatal on failure)
            download_url = None
            try:
                if connection_string and account_name and account_key and container_name and file_path:
                    sas_token = generate_blob_sas(
                        account_name=account_name,
                        container_name=container_name,
                        blob_name=file_path,
                        account_key=account_key,
                        permission=BlobSasPermissions(read=True),
                        expiry=datetime.utcnow() + timedelta(days=7),
                        content_disposition=f'attachment; filename="{os.path.basename(file_path)}"'
                    )
                    download_url = f"https://{account_name}.blob.{endpoint_suffix}/{container_name}/{file_path}?{sas_token}"
            except Exception:
                download_url = None

            tasks.append({
                "task_id": task_id,
                "status": status,
                "created_at": str(created_at),
                "updated_at": str(updated_at),
                "file": file_path,
                "file_name": os.path.basename(file_path) if file_path else None,  # <- added
                "domain": domain,
                "kb_name": kb_name,
                "file_size": file_size,
                "uploaded_by": uploaded_by,  # may be full name if resolved at insert time
                "container_name": container_name,
                "download_url": download_url,
            })

        result = {
            "workspace_id": workspace_id,
            "total": len(tasks),
            "tasks": tasks
        }
        if user_id:
            result["filtered_by_user"] = filter_uploaded_by

        if include_counts:
            counts_query = """
                SELECT status, COUNT(*)
                FROM file_tasks
                WHERE workspace_id = %s
            """
            count_params = [workspace_id]
            if filter_uploaded_by is not None:
                counts_query += " AND uploaded_by = %s"
                count_params.append(filter_uploaded_by)
            counts_query += " GROUP BY status"
            cur.execute(counts_query, tuple(count_params))
            counts = {status: count for status, count in cur.fetchall()}
            result["counts"] = counts

        cur.close()
        conn.close()

        if not tasks:
            result["message"] = "No indexing tasks found for the specified workspace_id."

        return result

    except Exception as e:
        return {"error": str(e)}


@mcp.tool()
async def generate_download_urls_by_workspace(
    ctx: Context,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
    limit: int = 1000,
    expiry_days: int = 7
) -> dict:
    """
    Generate SAS download URLs for all files in a workspace.
    - Pulls records from file_tasks for the given workspace_id
    - Optionally filters by uploaded_by derived from user_id
    - Returns a list of file metadata + SAS URLs
    Also returns a single ZIP SAS URL that bundles all existing files found.
    """
    if user_id is None:
        return {"status": "error", "error": "user_id cannot be null"}

    # --- JWT-based authentication and workspace-user mapping check (copied from list_workspace_users) ---
    if not workspace_id:
        return {"error": "workspace_id is required"}

    # Validate user access to workspace
    valid, err = validate_user_workspace_access(workspace_id=workspace_id)
    if not valid:
        await ctx.debug(f"User not mapped to workspace: {err}")
        return {"error": err}

    # Enforce JWT-based access: only allow if user is mapped to the workspace
    request = request_var.get(None)
    if not request or not hasattr(request.state, "jwt_claims"):
        await ctx.debug("Unauthorized: JWT claims not found in request context")
        return {"error": "Unauthorized: JWT claims not found in request context"}
    claims = request.state.jwt_claims
    jwt_user_id = claims.get("user_id") or claims.get("sub")
    if not jwt_user_id:
        await ctx.debug("Unauthorized: user_id not found in token claims")
        return {"error": "Unauthorized: user_id not found in token claims"}

    # Check if user is mapped to this workspace
    session = db.Session()
    try:
        user_map = session.query(db.UserMap).filter_by(workspace_id=workspace_id, user_id=jwt_user_id, is_active=True).first()
        if not user_map:
            session.close()
            await ctx.debug("You are not authorized to access this workspace.")
            return {"error": "You are not authorized to access this workspace."}
    except Exception as e:
        session.close()
        await ctx.debug(f"Error during workspace-user mapping check: {e}")
        return {"error": str(e)}
    finally:
        pass

    try:
        # Resolve optional user filter
        filter_uploaded_by = uploaded_by_username(str(user_id)) if user_id else None

        # Query file_tasks for the files in the workspace
        conn = psycopg2.connect(
            host=os.environ["POSTGRES_HOST"],
            user=os.environ["POSTGRES_USER"],
            password=os.environ["POSTGRES_PASSWORD"],
            dbname=os.environ["POSTGRESQL_DATABASE_DATABASE_2"]
        )
        cur = conn.cursor()

        base_query = """
            SELECT id, container_name, file_path, status, domain, kb_name, uploaded_by, created_at
            FROM file_tasks
            WHERE workspace_id = %s
        """
        params = [workspace_id]
        if filter_uploaded_by is not None:
            base_query += " AND uploaded_by = %s"
            params.append(filter_uploaded_by)
        base_query += " ORDER BY created_at DESC LIMIT %s"
        params.append(limit)

        cur.execute(base_query, tuple(params))
        rows = cur.fetchall()
        cur.close()
        conn.close()

        if not rows:
            return {"workspace_id": workspace_id, "files": [], "message": "No files found for this workspace."}

        # Prepare Azure client
        connection_string = os.getenv('AZURE_BLOB_STORAGE_CONNECTION_STRING')
        if not connection_string:
            return {"error": "AZURE_BLOB_STORAGE_CONNECTION_STRING is not set."}

        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        account_name = str(blob_service_client.account_name) if blob_service_client.account_name else ""
        if not account_name:
            return {"error": "Could not determine Azure Storage account name."}

        expiry = datetime.now() + timedelta(days=expiry_days)

        results = []
        files_for_zip = []  # (container_name, blob_path, file_name)
        first_container_name = None

        for (task_id, container_name, blob_path, status, domain, kb_name, uploaded_by, created_at) in rows:
            if not container_name or not blob_path:
                # skip incomplete records
                continue

            if first_container_name is None:
                first_container_name = container_name

            container_client = blob_service_client.get_container_client(container_name)
            blob_client = container_client.get_blob_client(blob_path)

            # Check existence to avoid dead links
            try:
                exists = blob_client.exists()
            except Exception:
                exists = False

            if not exists:
                results.append({
                    "task_id": task_id,
                    "domain": domain,
                    "kb_name": kb_name,
                    "file_path": blob_path,
                    "file_name": os.path.basename(blob_path),
                    "status": status,
                    "uploaded_by": uploaded_by,
                    "created_at": str(created_at),
                    "exists": False,
                    "download_url": None
                })
                continue

            # Generate read-only SAS URL
            try:
                sas_token = generate_blob_sas(
                    account_name=account_name,
                    container_name=container_name,
                    blob_name=blob_path,
                    account_key=blob_service_client.credential.account_key,
                    permission=BlobSasPermissions(read=True),
                    expiry=expiry,
                    # Force download in browser
                    content_disposition=f'attachment; filename="{os.path.basename(blob_path)}"'
                )
                download_url = f"https://{account_name}.blob.core.windows.net/{container_name}/{blob_path}?{sas_token}"
            except Exception as e:
                download_url = f"Error generating SAS: {str(e)}"

            # Track for ZIP
            files_for_zip.append((container_name, blob_path, os.path.basename(blob_path)))

            results.append({
                "task_id": task_id,
                "domain": domain,
                "kb_name": kb_name,
                "file_path": blob_path,
                "file_name": os.path.basename(blob_path),
                "status": status,
                "uploaded_by": uploaded_by,
                "created_at": str(created_at),
                "exists": True,
                "download_url": download_url
            })

        # Create a single ZIP containing all existing files (if any)
        zip_info = {
            "url": None,
            "container": None,
            "blob_path": None,
            "file_count": 0
        }

        if files_for_zip and first_container_name:
            try:
                zip_container_name = first_container_name
                zip_blob_path = f"workspace_zips/{workspace_id}/download_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"

                # Ensure zip container exists
                zip_container_client = blob_service_client.get_container_client(zip_container_name)
                try:
                    zip_container_client.get_container_properties()
                except Exception:
                    zip_container_client.create_container()

                # Build zip in a spooled temp file (spools to disk beyond threshold)
                spooled = tempfile.SpooledTemporaryFile(max_size=50 * 1024 * 1024, mode="w+b")
                with zipfile.ZipFile(spooled, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
                    for (cname, bpath, fname) in files_for_zip:
                        try:
                            cclient = blob_service_client.get_container_client(cname)
                            bclient = cclient.get_blob_client(bpath)
                            data = bclient.download_blob().readall()
                            zf.writestr(fname, data)
                        except Exception as e:
                            # Skip files that fail to download
                            continue

                spooled.seek(0)
                zip_blob_client = zip_container_client.get_blob_client(zip_blob_path)
                zip_blob_client.upload_blob(spooled, overwrite=True, content_settings=ContentSettings(content_type="application/zip"))
                spooled.close()

                # Generate SAS for ZIP
                zip_sas = generate_blob_sas(
                    account_name=account_name,
                    container_name=zip_container_name,
                    blob_name=zip_blob_path,
                    account_key=blob_service_client.credential.account_key,
                    permission=BlobSasPermissions(read=True),
                    expiry=expiry,
                    content_disposition=f'attachment; filename="{os.path.basename(zip_blob_path)}"'
                )
                zip_url = f"https://{account_name}.blob.core.windows.net/{zip_container_name}/{zip_blob_path}?{zip_sas}"

                zip_info = {
                    "url": zip_url,
                    "container": zip_container_name,
                    "blob_path": zip_blob_path,
                    "file_count": len(files_for_zip)
                }
            except Exception as e:
                zip_info = {
                    "url": f"Error creating ZIP: {str(e)}",
                    "container": first_container_name,
                    "blob_path": None,
                    "file_count": len(files_for_zip)
                }

        response = {
            "workspace_id": workspace_id,
            "total": len(results),
            "expiry_days": expiry_days,
            "files": results,
            "zip": zip_info
        }
        if user_id:
            response["filtered_by_user"] = filter_uploaded_by

        return response

    except Exception as e:
        return {"error": str(e)}
   
 
# @mcp.tool()
# async def conversation_indexing_tool(
#     ctx: Context,
#     domain: Optional[str] = None,
#     kb_name: Optional[str] = None,
#     url: Optional[str] = None
#     ) -> dict:
#     """
#     Index all user queries into the specified domain/KB directory.
#     """
#     try:
#         print("Starting crawling for URL...")
#         browser_config = BrowserConfig(verbose=True)
#         run_config = CrawlerRunConfig(
#                 markdown_generator=DefaultMarkdownGenerator(
#                 options={"ignore_links": True}
#             ),
#             # Content filtering
#             word_count_threshold=50,
#             excluded_tags=['form', 'header'],
#             exclude_external_links=True,
 
#             # Content processing
#             process_iframes=True,
#             remove_overlay_elements=True,
#         )
#         print("Browser and run config set up. Starting crawler...")
 
#         async with AsyncWebCrawler(config=browser_config) as crawler:
#             result = await crawler.arun(
#                 url=url,
#                 config=run_config
#             )
 
#         print("Crawling completed. Starting indexing...")
 
#         content = result.markdown
#         print("Content: ", content[:500])
 
#         def chunk_text(text, chunk_size=2000):
#             # Simple chunking by character count
#             return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
       
#         chunks = chunk_text(content, chunk_size=2000)
#         print("First chunk: ",chunks[0])
 
#         rag = await initialize_rag(domain=domain, kb_name=kb_name)
#         for idx, chunk in enumerate(chunks):
#             await rag.ainsert(input=chunk, file_paths=[url])
#             # await ctx.debug(f"Progress: {idx+1}/{len(chunks)}")
#         return {"status": "success", "file": url, "chunks": len(chunks)}
#     except Exception as e:
#         return {"error": str(e)}
    
@mcp.tool()
async def conversation_indexing_tool(
    ctx: Context,
    domain: Optional[str] = None,
    kb_name: Optional[str] = None,
    url: Optional[str] = None
    ) -> dict:
    """
    Index all user queries into the specified domain/KB directory.
    """
    try:
        print("Starting crawling for URL...")
        browser_config = BrowserConfig(verbose=True)
 
        # Do NOT drop link text on Wikipedia. Lower the word threshold.
        run_config = CrawlerRunConfig(
            markdown_generator=DefaultMarkdownGenerator(
                options={
                    "ignore_links": False  # keep anchor text
                }
            ),
            word_count_threshold=0,       # accept short sections
            excluded_tags=['form', 'header'],
            exclude_external_links=True,
            process_iframes=True,
            remove_overlay_elements=True,
        )
        print("Browser and run config set up. Starting crawler...")
 
        async with AsyncWebCrawler(config=browser_config) as crawler:
            result = await crawler.arun(url=url, config=run_config)
 
            content = (result.markdown or "").strip()
            # Fallback retry with minimal filtering if empty
            if not content:
                print("Empty markdown, retrying with minimal filtering...")
                retry_config = CrawlerRunConfig(
                    markdown_generator=DefaultMarkdownGenerator(options={"ignore_links": False}),
                    word_count_threshold=0
                )
                retry = await crawler.arun(url=url, config=retry_config)
                content = ((retry.markdown or "") or (getattr(retry, "cleaned_text", "") or "")).strip()
 
        print("Crawling completed. Starting indexing...")
        print("Content: ", content[:500])
 
        if not content:
            return {"error": "No content extracted from URL. Try a different URL or relax crawler filters."}
 
        def chunk_text(text, chunk_size=2000):
            return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
       
        chunks = chunk_text(content, chunk_size=2000)
        if not chunks:
            return {"error": "No chunks produced from content."}
        print("First chunk: ", chunks[0][:200])
 
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        for idx, chunk in enumerate(chunks):
            await rag.ainsert(input=chunk, file_paths=[url])
        return {"status": "success", "file": url, "chunks": len(chunks)}
    except Exception as e:
        return {"error": str(e)}    
 
# @mcp.tool()
# async def ingest_kb_file(domain: str = None, kb_name: str = None, file_name: str = None, file_bytes: bytes = None) -> str:
#     """
#     Save an uploaded file to the correct KB directory.
#     """
#     try:
#         if not domain or not kb_name or not file_name or not file_bytes:
#             raise ValueError("All parameters (domain, kb_name, file_name, file_bytes) are required.")
#         kb_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'data', domain, kb_name))
#         os.makedirs(kb_dir, exist_ok=True)
#         dest_path = os.path.join(kb_dir, file_name)
#         with open(dest_path, "wb") as f:
#             f.write(file_bytes)
#         return f"File {file_name} uploaded to {kb_dir}"
#     except Exception as e:
#         return f"Error: {e}"
   
# @mcp.tool()
# async def get_kb_knowledge_graph(
#     ctx: Context,
#     domain: Optional[str],
#     kb_name: Optional[str],
#     question: Optional[str],
#     role_id: Optional[int],
#     workspace_id: Optional[int],
#     history: Optional[List] = [],
#     mode: Optional[str] = "mix",
#     user_prompt: Optional[str] = "",
#     knowledge_bases: Optional[List[str]] = None
# ) -> dict:
#     """
#     Get the knowledge graph for the specified workspace (domain_kb_name) or for provided knowledge bases.
#     Returns nodes and edges as a dict for visualization.
#     """
#     if role_id != 34:
#         digit_map = {
#             '0': 'zero', '1': 'one', '2': 'two', '3': 'three', '4': 'four',
#             '5': 'five', '6': 'six', '7': 'seven', '8': 'eight', '9': 'nine'
#         }
#         result = []
#         for c in str(workspace_id):
#             if c.isalpha():
#                 result.append(c)
#             elif c.isdigit():
#                 result.append(digit_map[c])
#             # skip non-alphanumeric characters
#         workspace_id_alpha = ''.join(result)
#     #    kb_name = f"{kb_name}/{workspace_id_alpha}"

#     try:
#         all_nodes = []
#         all_edges = []
#         print("[DEBUG] get_kb_knowledge_graph called with:", {
#             "domain": domain,
#             "kb_name": kb_name,
#             "question": question,
#             "history": history,
#             "mode": mode,
#             "user_prompt": user_prompt,
#             "knowledge_bases": knowledge_bases
#         })
#         # If knowledge_bases are provided, query each and merge results
#         if knowledge_bases and isinstance(knowledge_bases, list) and len(knowledge_bases) > 0:
#             print(f"[DEBUG] knowledge_bases provided: {knowledge_bases}")
#             semaphore = asyncio.Semaphore(4)  # Lower concurrency to avoid event loop overload
#             async def fetch_graph_for_kb(kb):
#                 async with semaphore:
#                     try:
#                         combined_kb_name = f"{kb_name}{kb}"
#                         enhanced_question = f"{question} {domain} {kb_name}"
#                         print(f"[DEBUG] Initializing RAG for KB: domain={domain}, kb_name={combined_kb_name}")
#                         rag = await initialize_rag(domain=domain, kb_name=combined_kb_name)
#                         query_graph_data = await rag.aquery_data(
#                             enhanced_question, 
#                             question,
#                             param=QueryParam(
#                                 mode=mode,
#                                 top_k=5,
#                                 conversation_history=history,
#                                 user_prompt=user_prompt,
#                             )
#                         )
#                         print(f"[DEBUG] query_graph_data for KB '{kb}': {query_graph_data}")
#                         all_labels = query_graph_data["data"].get("entities", [])
#                         print(f"[DEBUG] all_labels for KB '{kb}': {all_labels}")
#                         if not all_labels:
#                             return kb, [], []
#                         label_semaphore = asyncio.Semaphore(4)
#                         async def fetch_graph(label):
#                             async with label_semaphore:
#                                 try:
#                                     print(f"[DEBUG] Fetching graph for label: {label}")
#                                     graph = await rag.get_knowledge_graph(label, 2, None)
#                                     print(f"[DEBUG] Graph for label '{label}': nodes={{}} edges={{}}".format(len(graph.nodes) if graph else 'None', len(graph.edges) if graph else 'None'))
#                                     return label, graph
#                                 except Exception as neo4j_err:
#                                     print(f"[DEBUG] Exception in fetch_graph for label '{label}': {neo4j_err}")
#                                     return label, None
#                         # Run fetch_graph in small batches to avoid timeouts
#                         batch_size = 4
#                         results = []
#                         label_names = [label['entity_name'] for label in all_labels]
#                         for i in range(0, len(label_names), batch_size):
#                             batch = label_names[i:i+batch_size]
#                             batch_results = await asyncio.gather(*(fetch_graph(label) for label in batch))
#                             results.extend(batch_results)
#                         nodes = []
#                         edges = []
#                         for idx, (label, graph) in enumerate(results):
#                             if graph:
#                                 nodes.extend([dict(node) for node in graph.nodes])
#                                 edges.extend([dict(edge) for edge in graph.edges])
#                         return kb, nodes, edges
#                     except Exception as neo4j_err:
#                         print(f"[DEBUG] Exception in fetch_graph_for_kb for KB '{kb}': {neo4j_err}")
#                         return kb, [], []
#             # Run fetch_graph_for_kb in small batches as well
#             batch_size_kb = 2
#             kb_results = []
#             for i in range(0, len(knowledge_bases), batch_size_kb):
#                 batch = knowledge_bases[i:i+batch_size_kb]
#                 batch_results = await asyncio.gather(*(fetch_graph_for_kb(kb) for kb in batch))
#                 kb_results.extend(batch_results)
#             for kb, nodes, edges in kb_results:
#                 all_nodes.extend(nodes)
#                 all_edges.extend(edges)
#             print(f"[DEBUG] Final merged nodes: {len(all_nodes)}, edges: {len(all_edges)}")
#             return {
#                 "knowledge_bases": knowledge_bases,
#                 "nodes": all_nodes,
#                 "edges": all_edges,
#             }
#         else:
#             if not domain or not kb_name:
#                 print("[DEBUG] Missing domain or kb_name")
#                 raise ValueError("Parameters 'domain', and 'kb_name' are required if knowledge_bases is not provided.")
#             print(f"[DEBUG] Initializing RAG for domain={domain}, kb_name={kb_name}")
#             rag = await initialize_rag(domain=domain, kb_name=kb_name)
#             workspace = f"{domain}_{kb_name}"
#             print(f"[DEBUG] Querying graph data with question: {question}")
#             query_graph_data = await rag.aquery_data(
#                 question,
#                 param=QueryParam(
#                     mode=mode,
#                     top_k=5,
#                     conversation_history=history,
#                     user_prompt=user_prompt,
#                 )
#             )
#             print(f"[DEBUG] query_graph_data: {query_graph_data}")
#             all_labels = query_graph_data["data"].get("entities", [])
#             print(f"[DEBUG] all_labels: {all_labels}")
#             if not all_labels:
#                 print(f"[DEBUG] No entities found for workspace {workspace}")
#                 return {workspace: {"nodes": [], "edges": []}}
#             semaphore = asyncio.Semaphore(4)
#             async def fetch_graph(label):
#                 async with semaphore:
#                     try:
#                         print(f"[DEBUG] Fetching graph for label: {label}")
#                         graph = await rag.get_knowledge_graph(label, 2, None)
#                         print(f"[DEBUG] Graph for label '{label}': nodes={len(graph.nodes) if graph else 'None'}, edges={len(graph.edges) if graph else 'None'}")
#                         return label, graph
#                     except Exception as neo4j_err:
#                         print(f"[DEBUG] Exception in fetch_graph for label '{label}': {neo4j_err}")
#                         return label, None
#             batch_size = 4
#             label_names = [label['entity_name'] for label in all_labels]
#             results = []
#             for i in range(0, len(label_names), batch_size):
#                 batch = label_names[i:i+batch_size]
#                 batch_results = await asyncio.gather(*(fetch_graph(label) for label in batch))
#                 results.extend(batch_results)
#             for idx, (label, graph) in enumerate(results):
#                 if graph:
#                     all_nodes.extend([dict(node) for node in graph.nodes])
#                     all_edges.extend([dict(edge) for edge in graph.edges])
#             print(f"[DEBUG] Final nodes: {len(all_nodes)}, edges: {len(all_edges)}")
#             return {
#                 workspace: {
#                     "nodes": all_nodes,
#                     "edges": all_edges,
#                 }
#             }
#     except Exception as e:
#         print(f"[DEBUG] get_kb_knowledge_graph failed: {e}")
#         return {"error": f"get_kb_knowledge_graph failed: {str(e)}"}

@mcp.tool()
async def get_kb_knowledge_graph(
    ctx: Context,
    domain: Optional[str],
    kb_name: Optional[str],
    question: Optional[str],
    role_id: Optional[int],
    workspace_id: Optional[int],
    history: Optional[List] = [],
    mode: Optional[str] = "mix",
    user_prompt: Optional[str] = "",
    knowledge_bases: Optional[List[str]] = []
) -> dict:
    """
    Get the knowledge graph for the specified workspace (domain_kb_name) or for provided knowledge bases.
    Returns nodes and edges as a dict for visualization with consistent format.
    """
    try:
        all_nodes = []
        all_edges = []
        digit_map = {
            '0': 'zero', '1': 'one', '2': 'two', '3': 'three', '4': 'four',
            '5': 'five', '6': 'six', '7': 'seven', '8': 'eight', '9': 'nine'
        }
        result = []
        for c in str(workspace_id):
            if c.isalpha():
                result.append(c)
            elif c.isdigit():
                result.append(digit_map[c])
        workspace_id_alpha = ''.join(result)
 
        knowledge_bases.append(workspace_id_alpha)
        print("[DEBUG] get_kb_knowledge_graph called with:", {
            "domain": domain,
            "kb_name": kb_name,
            "question": question,
            "knowledge_bases": knowledge_bases
        })
        
        # Validate required parameters
        if not question:
            return {
                "status": "error",
                "message": "Question parameter is required",
                "nodes": [],
                "edges": []
            }
        
        # If knowledge_bases are provided, query each and merge results
        if knowledge_bases and isinstance(knowledge_bases, list) and len(knowledge_bases) > 0:
            print(f"[DEBUG] Processing {len(knowledge_bases)} knowledge bases")

            # Initialize RAGs sequentially to avoid os.environ['NEO4J_DATABASE'] race condition
            # during parallel initialization. Query execution remains parallel below.
            rag_map = {}
            for kb in knowledge_bases:
                combined_kb_name = f"{kb_name}{kb}"
                print(f"[DEBUG] Initializing RAG for KB: domain={domain}, kb_name={combined_kb_name}")
                rag_map[kb] = await initialize_rag(domain=domain, kb_name=combined_kb_name)

            semaphore = asyncio.Semaphore(3)  # Reduced concurrency for stability

            async def fetch_graph_for_kb(kb):
                async with semaphore:
                    try:
                        rag = rag_map[kb]
                        #enhanced_question = f"{question} for {kb}"
                        enhanced_question = f"{question} {kb_name} {kb} "
                        print("Enhanced Question ",enhanced_question)
                        query_graph_data = await rag.aquery_data(
                            enhanced_question,  # Use enhanced question instead of original
                            param=QueryParam(
                                mode=mode,
                                top_k=5,
                                conversation_history=history,
                                user_prompt=user_prompt,
                            )
                        )
                        #print(f"[DEBUG] query_graph_data structure: {query_graph_data}")
                        
                        all_labels = query_graph_data.get("data", {}).get("entities", [])
                        print(f"[DEBUG] Found {len(all_labels)} entities for KB '{kb}'")
                        
                        if not all_labels:
                            return kb, [], []
                        
                        # Process labels with controlled parallelism
                        label_semaphore = asyncio.Semaphore(3)
                        
                        async def fetch_graph(label_info):
                            async with label_semaphore:
                                label = label_info.get('entity_name', '')
                                if not label:
                                    return None, None
                                try:
                                    print(f"[DEBUG] Fetching graph for label: {label}")
                                    graph = await rag.get_knowledge_graph(label, 2, None)
                                    if graph:
                                        print(f"[DEBUG] Graph for label '{label}': nodes={len(graph.nodes)}, edges={len(graph.edges)}")
                                        return graph.nodes, graph.edges
                                    return None, None
                                except Exception as e:
                                    print(f"[ERROR] Exception fetching graph for label '{label}': {e}")
                                    return None, None
                        
                        # Parallel fetch all labels for this KB
                        graph_results = await asyncio.gather(
                            *[fetch_graph(label) for label in all_labels],
                            return_exceptions=True
                        )
                        
                        nodes = []
                        edges = []
                        for result in graph_results:
                            if isinstance(result, tuple) and result[0] is not None:
                                node_list, edge_list = result
                                if node_list:
                                    nodes.extend([dict(node) for node in node_list])
                                if edge_list:
                                    edges.extend([dict(edge) for edge in edge_list])
                        
                        print(f"[DEBUG] KB '{kb}' produced {len(nodes)} nodes and {len(edges)} edges")
                        return kb, nodes, edges
                        
                    except Exception as e:
                        print(f"[ERROR] Exception processing KB '{kb}': {e}")
                        import traceback
                        traceback.print_exc()
                        return kb, [], []
            
            # Process all KBs in parallel with controlled concurrency
            print(f"[DEBUG] Starting parallel processing of {len(knowledge_bases)} KBs")
            kb_results = await asyncio.gather(
                *[fetch_graph_for_kb(kb) for kb in knowledge_bases],
                return_exceptions=True
            )
            
            # Merge all results
            for result in kb_results:
                if isinstance(result, tuple):
                    kb, nodes, edges = result
                    all_nodes.extend(nodes)
                    all_edges.extend(edges)
                elif isinstance(result, Exception):
                    print(f"[ERROR] KB processing failed: {result}")
            
            # Deduplicate nodes and edges
            unique_nodes = {node.get('id', str(i)): node for i, node in enumerate(all_nodes)}
            all_nodes = list(unique_nodes.values())
            
            print(f"[DEBUG] Final merged result: {len(all_nodes)} nodes, {len(all_edges)} edges")
            
            return {
                "status": "success",
                "knowledge_bases": knowledge_bases,
                "nodes": all_nodes,
                "edges": all_edges,
            }
        
        else:
            # Single workspace mode
            if not domain or not kb_name:
                return {
                    "status": "error",
                    "message": "Parameters 'domain' and 'kb_name' are required",
                    "nodes": [],
                    "edges": []
                }
            #  ENHANCE QUESTION WITH DOMAIN CONTEXT
            enhanced_question = f"{question} {domain} {kb_name} {knowledge_bases}"
            print(f"[DEBUG] Processing single workspace: {domain}_{kb_name}")
            rag = await initialize_rag(domain=domain, kb_name=kb_name)
            
            print("Enhanced Question ",enhanced_question)
            query_graph_data = await rag.aquery_data(
                enhanced_question, # Use enhanced question
                param=QueryParam(
                    mode=mode,
                    top_k=5,
                    conversation_history=history,
                    user_prompt=user_prompt,
                )
            )
            
            all_labels = query_graph_data.get("data", {}).get("entities", [])
            print(f"[DEBUG] Found {len(all_labels)} entities")
            
            if not all_labels:
                print(f"[DEBUG] No entities found")
                return {
                    "status": "success",
                    "message": "No entities found for the query",
                    "nodes": [],
                    "edges": []
                }
            
            semaphore = asyncio.Semaphore(3)
            
            async def fetch_graph(label_info):
                async with semaphore:
                    label = label_info.get('entity_name', '')
                    if not label:
                        return None, None
                    try:
                        print(f"[DEBUG] Fetching graph for label: {label}")
                        graph = await rag.get_knowledge_graph(label, 2, None)
                        if graph:
                            print(f"[DEBUG] Graph for label '{label}': nodes={len(graph.nodes)}, edges={len(graph.edges)}")
                            return graph.nodes, graph.edges
                        return None, None
                    except Exception as e:
                        print(f"[ERROR] Exception fetching graph for label '{label}': {e}")
                        return None, None
            
            # Parallel fetch all labels
            graph_results = await asyncio.gather(
                *[fetch_graph(label) for label in all_labels],
                return_exceptions=True
            )
            
            for result in graph_results:
                if isinstance(result, tuple) and result[0] is not None:
                    node_list, edge_list = result
                    if node_list:
                        all_nodes.extend([dict(node) for node in node_list])
                    if edge_list:
                        all_edges.extend([dict(edge) for edge in edge_list])
            
            # Deduplicate nodes
            unique_nodes = {node.get('id', str(i)): node for i, node in enumerate(all_nodes)}
            all_nodes = list(unique_nodes.values())
            
            print(f"[DEBUG] Final result: {len(all_nodes)} nodes, {len(all_edges)} edges")
            
            return {
                "status": "success",
                "nodes": all_nodes,
                "edges": all_edges,
            }
            
    except Exception as e:
        print(f"[ERROR] get_kb_knowledge_graph failed: {e}")
        import traceback
        traceback.print_exc()
        return {
            "status": "error",
            "message": f"Failed to fetch knowledge graph: {str(e)}",
            "nodes": [],
            "edges": []
        }
 
@mcp.tool()
async def insert_node_to_kg(
    domain: Optional[str] = None,
    kb_name: Optional[str] = None,
    entity_name: Optional[str] = None,
    entity_data: Optional[dict] = None
):
    try:
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        node_data = {
            "entity_id": entity_name,
            "entity_type": entity_data.get("entity_type", "UNKNOWN") if entity_data else "UNKNOWN",
            "description": entity_data.get("description", "") if entity_data else "",
            "source_id": entity_data.get("source_id", "manual_creation") if entity_data else "manual_creation",
            "file_path": entity_data.get("file_path", "manual_creation") if entity_data else "manual_creation",
            "created_at": int(time.time()),
        }
        return await rag.acreate_entity(
            entity_name=entity_name,
            entity_data=node_data
        )
    except Exception as e:
        return {"error": str(e)}
 
@mcp.tool()
async def insert_edge_to_kg(
    domain: Optional[str] = None,
    kb_name: Optional[str] = None,
    source_entity_name: Optional[str] = None,
    target_entity_name: Optional[str] = None,
    relation_data: Optional[dict] = None
):
    try:
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        edge_data = {
            "description": relation_data.get("description", "") if relation_data else "",
            "keywords": relation_data.get("keywords", "") if relation_data else "",
            "source_id": relation_data.get("source_id", "manual_creation") if relation_data else "manual_creation",
            "weight": float(relation_data.get("weight", 1.0)) if relation_data else 1.0,
            "file_path": relation_data.get("file_path", "manual_creation") if relation_data else "manual_creation",
            "created_at": int(time.time()),
        }
        return await rag.acreate_relation(
            source_entity=source_entity_name,
            target_entity=target_entity_name,
            relation_data=edge_data
        )
    except Exception as e:
        return {"error": str(e)}
 
@mcp.tool()
async def delete_entity_from_kg(
    domain: Optional[str] = None,
    kb_name: Optional[str] = None,
    entity_name: Optional[str] = None,
):
    try:
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        return await rag.adelete_by_entity(
            entity_name=entity_name,
        )
    except Exception as e:
        return {"error": str(e)}
   
@mcp.tool()
async def delete_relation_from_kg(
    domain: Optional[str] = None,
    kb_name: Optional[str] = None,
    source_entity_name: Optional[str] = None,
    target_entity_name: Optional[str] = None,
):
    try:
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        return await rag.adelete_by_relation(
            source_entity=source_entity_name,
            target_entity=target_entity_name,
        )
    except Exception as e:
        return {"error": str(e)}
 
@mcp.tool()
async def edit_entity_in_kg(
    domain: Optional[str] = None,
    kb_name: Optional[str] = None,
    entity_name: Optional[str] = None,
    updated_data: Optional[dict] = None,
):
    try:
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        return await rag.aedit_entity(
            entity_name=entity_name,
            updated_data=updated_data,
            allow_rename = True
        )
    except Exception as e:
        return {"error": str(e)}
 
@mcp.tool()
async def edit_relation_in_kg(
    domain: Optional[str] = None,
    kb_name: Optional[str] = None,
    source_entity_name: Optional[str] = None,
    target_entity_name: Optional[str] = None,
    updated_data: Optional[dict] = None,
):
    try:
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        return await rag.aedit_relation(
            source_entity=source_entity_name,
            target_entity=target_entity_name,
            updated_data=updated_data,
        )
    except Exception as e:
        return {"error": str(e)}
   
@mcp.tool()
async def extract_keywords_from_query(
    user_query: Optional[str] = None,
    history: Optional[list] = None,
    node_labels: Optional[list] = None
) -> dict:
    """
    Extract keywords from a user query, considering node labels as context.
    Returns only keywords related to the user prompt and the provided node labels.
    """
    try:
        if not user_query or not isinstance(user_query, str):
            return {"error": "user_query must be a non-empty string."}
        node_labels = node_labels or []
        node_labels_str = ', '.join([str(label) for label in node_labels]) if node_labels else ''
        system_prompt = (
            f"""Given the following node labels in the knowledge graph: [{node_labels_str}].
            Extract node labels from the user query that are most relevant to these node labels.
            Return only these node labels that match or are closely related to the node labels as strings in a list."""
        )
        keywords = await llm_model_func(prompt=user_query, system_prompt=system_prompt)
        return {"keywords": keywords}
    except Exception as e:
        return {"error": str(e)}
   
@mcp.tool()
async def get_indexed_file_names(domain, kb_name):
    """
    Get indexed file names and their doc IDs from shared doc status storage.
    This is resilient in multi-instance cloud deployments.
    """
    try:
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        doc_dict: Dict[str, List[str]] = {}
        workspace_name = ''.join(char for char in f"{domain}{kb_name}" if char.isalpha())

        page = 1
        page_size = 500

        while True:
            docs_page, total_count = await rag.doc_status.get_docs_paginated(
                status_filter=None,
                page=page,
                page_size=page_size,
                sort_field="updated_at",
                sort_direction="desc",
            )

            if not docs_page:
                break

            for doc_id, doc_status in docs_page:
                file_path = getattr(doc_status, "file_path", "") if doc_status else ""
                file_name = os.path.basename(file_path) if file_path else ""
                if not file_name:
                    continue

                if file_name not in doc_dict:
                    doc_dict[file_name] = []
                if doc_id not in doc_dict[file_name]:
                    doc_dict[file_name].append(doc_id)

            if page * page_size >= int(total_count or 0):
                break
            page += 1

        # Fallback path: some environments can have searchable chunks but empty doc_status.
        # In that case, derive file/doc mapping from chunk storage table directly.
        if not doc_dict:
            try:
                conn = psycopg2.connect(
                    host=os.environ["POSTGRES_HOST"],
                    user=os.environ["POSTGRES_USER"],
                    password=os.environ["POSTGRES_PASSWORD"],
                    dbname=os.environ.get("POSTGRES_DATABASE") or os.environ.get("POSTGRESQL_DATABASE_DATABASE_2")
                )
                cur = conn.cursor()
                cur.execute(
                    """
                    SELECT DISTINCT full_doc_id, file_path
                    FROM LIGHTRAG_VDB_CHUNKS
                    WHERE workspace = %s
                      AND full_doc_id IS NOT NULL
                      AND file_path IS NOT NULL
                    """,
                    (workspace_name,),
                )
                rows = cur.fetchall()
                cur.close()
                conn.close()

                for full_doc_id, file_path in rows:
                    file_name = os.path.basename(file_path) if file_path else ""
                    if not file_name or not full_doc_id:
                        continue
                    if file_name not in doc_dict:
                        doc_dict[file_name] = []
                    if full_doc_id not in doc_dict[file_name]:
                        doc_dict[file_name].append(full_doc_id)

                print(f"Fallback doc-chunks mapping used for workspace: {workspace_name}")
            except Exception as fallback_err:
                print(f"Fallback from LIGHTRAG_DOC_CHUNKS failed: {fallback_err}")

        print("Doc ID to File Name Dictionary:", doc_dict)
        return doc_dict
    except Exception as e:
        print(f"Failed to get indexed file names: {e}")
        return {}
 
#@mcp.tool()
# async def delete_by_doc_id(doc_id):
#     """
#     Delete a document by doc_id using LightRAG
#     """
#     try:
#         rag = await initialize_rag()
#         response = await rag.adelete_by_doc_id(doc_id)
#         await rag.aclear_cache()
#         print(f"Document with doc_id '{doc_id}' deleted successfully.")
#         return response
#     except Exception as e:
#         print(f"Error deleting document with doc_id '{doc_id}': {e}")

@mcp.tool()
async def delete_by_doc_id(doc_id, domain: Optional[str] = None, kb_name: Optional[str] = None):
    """
    Delete a document by doc_id using LightRAG
    """
    try:
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        response = await rag.adelete_by_doc_id(doc_id)
        await rag.aclear_cache()
        print(f"Document with doc_id '{doc_id}' deleted successfully.")
        return response
    except Exception as e:
        error_msg = f"Failed to delete document with doc_id '{doc_id}'. Please Verify the document ID exists and try again."
        return {"error": error_msg}  # Return error instead of None

def _delete_orphaned_vdb_chunks(doc_id: str, workspace_name: str) -> int:
    """
    Directly remove orphaned rows from LIGHTRAG_VDB_CHUNKS when LightRAG KV store
    no longer holds the doc (adelete_by_doc_id returns not_found).
    Returns total rows deleted across the table.
    """
    try:
        conn = psycopg2.connect(
            host=os.environ["POSTGRES_HOST"],
            user=os.environ["POSTGRES_USER"],
            password=os.environ["POSTGRES_PASSWORD"],
            dbname=os.environ.get("POSTGRES_DATABASE") or os.environ.get("POSTGRESQL_DATABASE_DATABASE_2"),
        )
        cur = conn.cursor()
        cur.execute(
            "DELETE FROM LIGHTRAG_VDB_CHUNKS WHERE full_doc_id = %s AND workspace = %s",
            (doc_id, workspace_name),
        )
        deleted = cur.rowcount
        conn.commit()
        cur.close()
        conn.close()
        return deleted
    except Exception as e:
        print(f"_delete_orphaned_vdb_chunks error doc_id={doc_id} workspace={workspace_name}: {e}")
        return 0


@mcp.tool()
async def delete_by_doc_ids(
    doc_ids: List[str],
    domain: Optional[str] = None,
    kb_name: Optional[str] = None,
    skip_rebuild: bool = True,
    trace_id: Optional[str] = None,
):
    """
    Bulk delete documents by doc_id using one initialized LightRAG session.
    """
    if not isinstance(doc_ids, list) or not doc_ids:
        return {
            "status": "error",
            "message": "doc_ids must be a non-empty list",
            "summary": {"requested": 0, "success": 0, "not_found": 0, "failed": 0},
            "results": []
        }

    rag = None
    original_rebuild_func = None
    results = []
    success_count = 0
    not_found_count = 0
    failed_count = 0

    print(
        f"DELETE_TRACE [{trace_id or 'n/a'}] delete_by_doc_ids:start "
        f"domain={domain} kb_name={kb_name} doc_count={len(doc_ids)} skip_rebuild={skip_rebuild}"
    )

    try:
        rag = await initialize_rag(domain=domain, kb_name=kb_name)

        # Fast file-purge mode: skip expensive per-doc KG rebuild cycle.
        if skip_rebuild:
            original_rebuild_func = getattr(lightrag_core, "rebuild_knowledge_from_chunks", None)

            async def _noop_rebuild_knowledge_from_chunks(*args, **kwargs):
                print("Fast purge mode enabled: skipping rebuild_knowledge_from_chunks")
                return None

            if original_rebuild_func is not None:
                lightrag_core.rebuild_knowledge_from_chunks = _noop_rebuild_knowledge_from_chunks

        for doc_id in doc_ids:
            try:
                print(
                    f"DELETE_TRACE [{trace_id or 'n/a'}] delete_by_doc_ids:doc_start "
                    f"kb_name={kb_name} doc_id={doc_id}"
                )
                # Hard-delete mode: also remove LLM cache tied to this doc to prevent KG rebuild from cache.
                response = await rag.adelete_by_doc_id(doc_id, delete_llm_cache=True)

                status = None
                status_code = None
                if isinstance(response, dict):
                    status = response.get("status")
                    status_code = response.get("status_code")
                    response_message = str(response.get("message") or "")
                else:
                    # LightRAG commonly returns a DeletionResult object
                    status = getattr(response, "status", None)
                    status_code = getattr(response, "status_code", None)
                    response_message = str(getattr(response, "message", "") or "")

                if status_code is None and status == "success":
                    status_code = 200
                if status_code is None and status == "not_found":
                    status_code = 404

                if status_code == 404 or status == "not_found":
                    # LightRAG KV store missing this doc — clean up orphaned VDB chunks directly.
                    workspace_name = ''.join(c for c in f"{domain or ''}{kb_name or ''}" if c.isalpha())
                    vdb_deleted = _delete_orphaned_vdb_chunks(doc_id, workspace_name)
                    if vdb_deleted:
                        print(
                            f"DELETE_TRACE [{trace_id or 'n/a'}] delete_by_doc_ids:orphan_cleanup "
                            f"kb_name={kb_name} doc_id={doc_id} vdb_rows_deleted={vdb_deleted}"
                        )
                        success_count += 1
                        status = "success"
                        status_code = 200
                    else:
                        not_found_count += 1
                elif status_code in (200, 204) or status == "success":
                    success_count += 1
                else:
                    # Treat non-success outcomes (for example 403/not_allowed) as failures.
                    failed_count += 1

                results.append({
                    "doc_id": doc_id,
                    "status": status or ("success" if status_code == 200 else "unknown"),
                    "status_code": status_code,
                    "response": response,
                })

                print(
                    f"DELETE_TRACE [{trace_id or 'n/a'}] delete_by_doc_ids:doc_result "
                    f"kb_name={kb_name} doc_id={doc_id} status={status} status_code={status_code} "
                    f"message={response_message[:300]}"
                )
            except Exception as e:
                failed_count += 1
                results.append({
                    "doc_id": doc_id,
                    "status": "failed",
                    "status_code": 500,
                    "message": str(e),
                })
                print(
                    f"DELETE_TRACE [{trace_id or 'n/a'}] delete_by_doc_ids:doc_exception "
                    f"kb_name={kb_name} doc_id={doc_id} error={str(e)}"
                )

        summary = {
            "requested": len(doc_ids),
            "success": success_count,
            "not_found": not_found_count,
            "failed": failed_count,
        }

        print(
            f"DELETE_TRACE [{trace_id or 'n/a'}] delete_by_doc_ids:summary "
            f"domain={domain} kb_name={kb_name} summary={summary}"
        )
        return {
            "status": "completed",
            "summary": summary,
            "results": results,
        }
    except Exception as e:
        print(
            f"DELETE_TRACE [{trace_id or 'n/a'}] delete_by_doc_ids:error "
            f"domain={domain} kb_name={kb_name} error={str(e)}"
        )
        return {
            "status": "error",
            "message": f"Bulk delete failed: {str(e)}",
            "summary": {
                "requested": len(doc_ids),
                "success": success_count,
                "not_found": not_found_count,
                "failed": failed_count + len(doc_ids) - (success_count + not_found_count + failed_count),
            },
            "results": results,
        }
    finally:
        if original_rebuild_func is not None:
            lightrag_core.rebuild_knowledge_from_chunks = original_rebuild_func

        if rag is not None:
            try:
                await rag.aclear_cache()
                print(f"DELETE_TRACE [{trace_id or 'n/a'}] delete_by_doc_ids:cache_cleared kb_name={kb_name}")
            except Exception as e:
                print(
                    f"DELETE_TRACE [{trace_id or 'n/a'}] delete_by_doc_ids:cache_clear_error "
                    f"kb_name={kb_name} error={str(e)}"
                )


@mcp.tool()
async def delete_file_single_call(
    domain: str,
    kb_name: str,
    file_name: str,
    workspace_id: Optional[str] = None,
    role_id: Optional[int] = None,
    skip_rebuild: bool = True,
):
    """
    One-call delete flow for frontend button:
    1) Resolve doc_ids for file_name across candidate KB scopes.
    2) Delete indexed records by doc_ids (per scope).
    3) Delete blob file from candidate containers/paths.
    """

    try:
        trace_id = f"del-{uuid.uuid4().hex[:12]}"
        print(
            f"DELETE_TRACE [{trace_id}] delete_file_single_call:start "
            f"domain={domain} kb_name={kb_name} file_name={file_name} "
            f"workspace_id={workspace_id} role_id={role_id} skip_rebuild={skip_rebuild}"
        )

        if not domain or not kb_name or not file_name:
            print(
                f"DELETE_TRACE [{trace_id}] delete_file_single_call:invalid_input "
                f"domain={domain} kb_name={kb_name} file_name={file_name}"
            )
            return {
                "status": "error",
                "message": "domain, kb_name and file_name are required",
                "resolved": {"doc_ids": []},
                "index_delete": {"scopes": [], "summary": {"requested": 0, "success": 0, "not_found": 0, "failed": 0}},
                "blob_delete": {"deleted": []},
            }

        # Build candidate KB scopes to tolerate role/type drift and legacy paths.
        workspace_id_alpha = ""
        if workspace_id:
            digit_map = {
                '0': 'zero', '1': 'one', '2': 'two', '3': 'three', '4': 'four',
                '5': 'five', '6': 'six', '7': 'seven', '8': 'eight', '9': 'nine'
            }
            result = []
            for c in str(workspace_id):
                if c.isalpha():
                    result.append(c)
                elif c.isdigit():
                    result.append(digit_map[c])
            workspace_id_alpha = ''.join(result)

        candidate_kbs = []
        if workspace_id_alpha:
            candidate_kbs.append(f"{kb_name}/{workspace_id_alpha}")
        if workspace_id:
            candidate_kbs.append(f"{kb_name}/{workspace_id}")
        candidate_kbs.append(kb_name)

        seen_kbs = set()
        candidate_kbs = [k for k in candidate_kbs if not (k in seen_kbs or seen_kbs.add(k))]

        print(
            f"DELETE_TRACE [{trace_id}] delete_file_single_call:candidate_kbs "
            f"candidate_kbs={candidate_kbs}"
        )

        base_filename = os.path.basename(file_name).strip()
        max_attempts = 20
        retry_delay_seconds = 1.2
        min_attempts_when_no_index_evidence = 10

        async def _resolve_doc_ids_once() -> tuple[list[str], dict[str, list[str]]]:
            doc_ids_by_kb_local: Dict[str, List[str]] = {}
            all_doc_ids_local: List[str] = []

            # First: resolve from shared doc status.
            for kb_scope in candidate_kbs:
                scope_doc_ids = []
                try:
                    rag_scope = await initialize_rag(domain=domain, kb_name=kb_scope)
                    page = 1
                    page_size = 500

                    while True:
                        docs_page, total_count = await rag_scope.doc_status.get_docs_paginated(
                            status_filter=None,
                            page=page,
                            page_size=page_size,
                            sort_field="updated_at",
                            sort_direction="desc",
                        )
                        print("0"*60)
                        print(docs_page)
                        print(total_count)
                        print("0"*60)
                        if not docs_page:
                            break

                        for doc_id, doc_status in docs_page:
                            file_path = getattr(doc_status, "file_path", "") if doc_status else ""
                            file_path_norm = str(file_path).replace("\\", "/")
                            file_name_from_status = file_path_norm.split("/")[-1].strip() if file_path_norm else ""
                            if file_name_from_status and file_name_from_status.lower() == base_filename.lower():
                                if doc_id and doc_id not in scope_doc_ids:
                                    scope_doc_ids.append(doc_id)
                                if doc_id and doc_id not in all_doc_ids_local:
                                    all_doc_ids_local.append(doc_id)

                        if page * page_size >= int(total_count or 0):
                            break
                        page += 1
                except Exception as status_err:
                    print(
                        f"DELETE_TRACE [{trace_id}] resolve_doc_ids:doc_status_error "
                        f"kb_scope={kb_scope} error={str(status_err)}"
                    )

                if scope_doc_ids:
                    doc_ids_by_kb_local[kb_scope] = scope_doc_ids

                print(
                    f"DELETE_TRACE [{trace_id}] resolve_doc_ids:doc_status_scan "
                    f"kb_scope={kb_scope} matched_doc_ids={len(scope_doc_ids)}"
                )

            # Fallback: resolve directly from chunk table for scopes still missing.
            try:
                conn = psycopg2.connect(
                    host=os.environ["POSTGRES_HOST"],
                    user=os.environ["POSTGRES_USER"],
                    password=os.environ["POSTGRES_PASSWORD"],
                    dbname=os.environ.get("POSTGRES_DATABASE") or os.environ.get("POSTGRESQL_DATABASE_DATABASE_2")
                )
                cur = conn.cursor()

                for kb_scope in candidate_kbs:
                    if kb_scope in doc_ids_by_kb_local and doc_ids_by_kb_local[kb_scope]:
                        continue

                    workspace = ''.join(char for char in f"{domain}{kb_scope}" if char.isalpha())
                    cur.execute(
                        """
                        SELECT DISTINCT full_doc_id
                        FROM LIGHTRAG_VDB_CHUNKS
                        WHERE workspace = %s
                          AND full_doc_id IS NOT NULL
                          AND file_path IS NOT NULL
                          AND (
                                lower(file_path) = lower(%s)
                                OR lower(file_path) LIKE lower(%s)
                                OR lower(file_path) LIKE lower(%s)
                                OR lower(file_path) LIKE lower(%s)
                          )
                        """,
                        (
                            workspace,
                            base_filename,
                            f"%/{base_filename}",
                            f"%\\{base_filename}",
                            f"%{base_filename}%",
                        ),
                    )
                    rows = cur.fetchall()
                    if not rows:
                        continue

                    scope_doc_ids = []
                    for (doc_id,) in rows:
                        if doc_id and doc_id not in scope_doc_ids:
                            scope_doc_ids.append(doc_id)
                        if doc_id and doc_id not in all_doc_ids_local:
                            all_doc_ids_local.append(doc_id)

                    if scope_doc_ids:
                        doc_ids_by_kb_local[kb_scope] = scope_doc_ids
                        print(
                            f"DELETE_TRACE [{trace_id}] resolve_doc_ids:fallback_match "
                            f"kb_scope={kb_scope} fallback_doc_ids={len(scope_doc_ids)}"
                        )

                cur.close()
                conn.close()
            except Exception as db_resolve_err:
                print(
                    f"DELETE_TRACE [{trace_id}] resolve_doc_ids:fallback_error "
                    f"error={str(db_resolve_err)}"
                )

            print(
                f"DELETE_TRACE [{trace_id}] resolve_doc_ids:summary "
                f"total_unique_doc_ids={len(all_doc_ids_local)} "
                f"scopes_with_docs={list(doc_ids_by_kb_local.keys())}"
            )

            return all_doc_ids_local, doc_ids_by_kb_local

        async def _delete_indexes_once(doc_ids_map: Dict[str, List[str]]) -> tuple[list[dict], int, int, int, int]:
            scope_results = []
            requested = 0
            success = 0
            not_found = 0
            failed = 0

            for kb_scope, scope_doc_ids in doc_ids_map.items():
                requested += len(scope_doc_ids)
                print(
                    f"DELETE_TRACE [{trace_id}] index_delete:scope_start "
                    f"kb_scope={kb_scope} doc_count={len(scope_doc_ids)}"
                )
                delete_result = await delete_by_doc_ids(
                    doc_ids=scope_doc_ids,
                    domain=domain,
                    kb_name=kb_scope,
                    skip_rebuild=skip_rebuild,
                    trace_id=trace_id,
                )

                summary = delete_result.get("summary", {}) if isinstance(delete_result, dict) else {}
                success += int(summary.get("success", 0) or 0)
                not_found += int(summary.get("not_found", 0) or 0)
                failed += int(summary.get("failed", 0) or 0)

                scope_results.append(
                    {
                        "kb_name": kb_scope,
                        "doc_ids": scope_doc_ids,
                        "result": delete_result,
                    }
                )

                print(
                    f"DELETE_TRACE [{trace_id}] index_delete:scope_result "
                    f"kb_scope={kb_scope} summary={summary}"
                )

            return scope_results, requested, success, not_found, failed

        connection_string = os.getenv('AZURE_BLOB_STORAGE_CONNECTION_STRING')
        main_container = os.getenv('AZURE_BLOB_STORAGE_CONTAINER_NAME')
        container_candidates = [c for c in [main_container, "workspace"] if c]
        blob_service_client = BlobServiceClient.from_connection_string(connection_string) if connection_string else None

        def _delete_blob_once() -> list[str]:
            if not blob_service_client:
                return []

            deleted_now = []
            for container_name in container_candidates:
                container_client = blob_service_client.get_container_client(container_name)
                for kb_scope in candidate_kbs:
                    blob_path = f"{domain}/{kb_scope}/{base_filename}"
                    blob_client = container_client.get_blob_client(blob_path)
                    try:
                        blob_client.delete_blob()
                        deleted_now.append(f"{container_name}:{blob_path}")
                        print(
                            f"DELETE_TRACE [{trace_id}] blob_delete:deleted "
                            f"container={container_name} blob_path={blob_path}"
                        )
                    except Exception:
                        # Missing blob in a candidate path/container is expected; continue.
                        pass
            return deleted_now

        def _blob_exists() -> bool:
            if not blob_service_client:
                return False

            for container_name in container_candidates:
                container_client = blob_service_client.get_container_client(container_name)
                for kb_scope in candidate_kbs:
                    blob_path = f"{domain}/{kb_scope}/{base_filename}"
                    blob_client = container_client.get_blob_client(blob_path)
                    try:
                        if blob_client.exists():
                            return True
                    except Exception:
                        continue
            return False

        index_scope_results = []
        total_requested = 0
        total_success = 0
        total_not_found = 0
        total_failed = 0
        resolved_doc_ids_all: List[str] = []
        resolved_by_kb_final: Dict[str, List[str]] = {}
        deleted_blob_paths_set = set()
        attempts_used = 0

        for attempt in range(1, max_attempts + 1):
            attempts_used = attempt
            print(
                f"DELETE_TRACE [{trace_id}] delete_file_single_call:attempt_start "
                f"attempt={attempt}/{max_attempts} file_name={base_filename}"
            )

            resolved_doc_ids, resolved_by_kb = await _resolve_doc_ids_once()
            print(
                f"DELETE_TRACE [{trace_id}] delete_file_single_call:attempt_resolved "
                f"attempt={attempt} resolved_doc_ids={len(resolved_doc_ids)} "
                f"scopes={list(resolved_by_kb.keys())}"
            )
            for doc_id in resolved_doc_ids:
                if doc_id not in resolved_doc_ids_all:
                    resolved_doc_ids_all.append(doc_id)
            for kb_scope, ids in resolved_by_kb.items():
                if kb_scope not in resolved_by_kb_final:
                    resolved_by_kb_final[kb_scope] = []
                for doc_id in ids:
                    if doc_id not in resolved_by_kb_final[kb_scope]:
                        resolved_by_kb_final[kb_scope].append(doc_id)

            if resolved_by_kb:
                scope_results, requested, success, not_found, failed = await _delete_indexes_once(resolved_by_kb)
                index_scope_results.extend(scope_results)
                total_requested += requested
                total_success += success
                total_not_found += not_found
                total_failed += failed

            deleted_now = _delete_blob_once()
            for path in deleted_now:
                deleted_blob_paths_set.add(path)

            # Verify both sides after the deletion pass.
            unresolved_doc_ids_after, _ = await _resolve_doc_ids_once()
            blob_still_exists = _blob_exists()
            print(
                f"DELETE_TRACE [{trace_id}] delete_file_single_call:attempt_verify "
                f"attempt={attempt} unresolved_doc_ids={len(unresolved_doc_ids_after)} "
                f"blob_still_exists={blob_still_exists}"
            )
            if not unresolved_doc_ids_after and not blob_still_exists:
                index_delete_evidence = (total_requested > 0) or (len(resolved_doc_ids_all) > 0)
                if index_delete_evidence:
                    print(
                        f"DELETE_TRACE [{trace_id}] delete_file_single_call:completed_early "
                        f"attempt={attempt} index_delete_evidence={index_delete_evidence}"
                    )
                    break

                if attempt >= min_attempts_when_no_index_evidence:
                    print(
                        f"DELETE_TRACE [{trace_id}] delete_file_single_call:stop_without_index_evidence "
                        f"attempt={attempt} min_attempts_when_no_index_evidence={min_attempts_when_no_index_evidence}"
                    )
                    break

            if attempt < max_attempts:
                print(
                    f"DELETE_TRACE [{trace_id}] delete_file_single_call:retry_sleep "
                    f"attempt={attempt} sleep_seconds={retry_delay_seconds}"
                )
                await asyncio.sleep(retry_delay_seconds)

        final_unresolved_doc_ids, _ = await _resolve_doc_ids_once()
        final_blob_exists = _blob_exists()

        index_delete_evidence_final = (total_requested > 0) or (len(resolved_doc_ids_all) > 0)
        completed_both = (not final_unresolved_doc_ids) and (not final_blob_exists) and index_delete_evidence_final
        status = "completed" if completed_both else "partial"

        print(
            f"DELETE_TRACE [{trace_id}] delete_file_single_call:final "
            f"status={status} attempts_used={attempts_used} "
            f"resolved_total_doc_ids={len(resolved_doc_ids_all)} "
            f"remaining_doc_ids={len(final_unresolved_doc_ids)} "
            f"blob_remaining_exists={final_blob_exists} "
            f"index_delete_evidence={index_delete_evidence_final} "
            f"index_summary={{'requested': {total_requested}, 'success': {total_success}, 'not_found': {total_not_found}, 'failed': {total_failed}}}"
        )

        if completed_both:
            final_message = "Completed deletion across index and blob"
        elif (not final_blob_exists) and (not index_delete_evidence_final):
            final_message = "Blob deleted but index document IDs were never resolved; index cleanup pending"
        else:
            final_message = "Reached retry limit before full deletion"

        return {
            "status": status,
            "trace_id": trace_id,
            "file_name": base_filename,
            "candidate_kbs": candidate_kbs,
            "attempts": {
                "used": attempts_used,
                "max": max_attempts,
            },
            "resolved": {
                "doc_ids": resolved_doc_ids_all,
                "by_kb": resolved_by_kb_final,
                "remaining_doc_ids": final_unresolved_doc_ids,
            },
            "index_delete": {
                "scopes": index_scope_results,
                "summary": {
                    "requested": total_requested,
                    "success": total_success,
                    "not_found": total_not_found,
                    "failed": total_failed,
                },
            },
            "blob_delete": {
                "deleted": sorted(list(deleted_blob_paths_set)),
                "remaining_exists": final_blob_exists,
            },
            "message": final_message,
        }
    except Exception as e:
        print(
            f"DELETE_TRACE [n/a] delete_file_single_call:error "
            f"domain={domain} kb_name={kb_name} file_name={file_name} error={str(e)}"
        )
        return {
            "status": "error",
            "message": f"delete_file_single_call failed: {str(e)}",
            "resolved": {"doc_ids": []},
            "index_delete": {"scopes": [], "summary": {"requested": 0, "success": 0, "not_found": 0, "failed": 0}},
            "blob_delete": {"deleted": []},
        }
 
@mcp.tool()
async def clear_cache(doc_id):
     try:
        rag = await initialize_rag()
        await rag.aclear_cache()
     except Exception as e:
        print(f"Error clearing cache': {e}")